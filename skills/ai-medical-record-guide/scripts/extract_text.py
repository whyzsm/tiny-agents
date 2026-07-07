#!/usr/bin/env python3
"""
病历文件文本提取工具

从 .pdf / .docx / .txt 文件中提取纯文本内容，供 AI 智能病历 skill 解析使用。

使用方式：
    python3 extract_text.py <file_path>

支持格式：
    - .txt：直接读取
    - .pdf：使用 PyPDF2 / pdfplumber 提取（需安装依赖）
    - .docx：使用 python-docx 提取（需安装依赖）

输出：
    提取的纯文本内容输出到 stdout。
"""

import sys
import os


def extract_from_txt(file_path: str) -> str:
    """从 .txt 文件读取文本"""
    encodings = ["utf-8", "gbk", "gb2312", "gb18030", "big5", "latin-1"]
    for enc in encodings:
        try:
            with open(file_path, "r", encoding=enc) as f:
                return f.read()
        except (UnicodeDecodeError, UnicodeError):
            continue
    raise ValueError(f"无法识别文件编码: {file_path}")


def extract_from_pdf(file_path: str) -> str:
    """从 .pdf 文件提取文本"""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        if text_parts:
            return "\n\n".join(text_parts)
    except ImportError:
        pass

    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        if text_parts:
            return "\n\n".join(text_parts)
    except ImportError:
        pass

    raise ImportError(
        "需要安装 PDF 处理库。请运行：pip3 install pdfplumber 或 pip3 install PyPDF2"
    )


def extract_from_docx(file_path: str) -> str:
    """从 .docx 文件提取文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    text_parts.append(row_text)

        return "\n".join(text_parts)
    except ImportError:
        raise ImportError(
            "需要安装 python-docx 库。请运行：pip3 install python-docx"
        )


def extract_text(file_path: str) -> str:
    """根据文件扩展名选择提取方式"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return extract_from_txt(file_path)
    elif ext == ".pdf":
        return extract_from_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return extract_from_docx(file_path)
    else:
        # 尝试作为纯文本读取
        try:
            return extract_from_txt(file_path)
        except Exception:
            raise ValueError(f"不支持的文件格式: {ext}")


def main():
    if len(sys.argv) < 2:
        print("用法: python3 extract_text.py <file_path>", file=sys.stderr)
        print("支持格式: .txt, .pdf, .docx", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        text = extract_text(file_path)
        print(text)
    except FileNotFoundError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)
    except ImportError as e:
        print(f"依赖缺失: {e}", file=sys.stderr)
        sys.exit(1)
    except ValueError as e:
        print(f"错误: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
