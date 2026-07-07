#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能老师 - 试卷生成脚本

目标：
1. 至少为九大学科生成可用的基础/提高/挑战三层练习。
2. 输入或依赖异常时给中文可操作提示，避免直接抛出英文堆栈。
3. 按 --output 真实写入文件；优先生成 docx，依赖缺失时降级为 Markdown。
"""

import argparse
import json
import sys
from pathlib import Path


SUPPORTED_SUBJECTS = ["语文", "数学", "英语", "物理", "化学", "生物", "历史", "地理", "政治"]
SKILL_ROOT = Path(__file__).resolve().parents[1]
BALANCED_BANK_PATH = SKILL_ROOT / "references" / "balanced_topic_bank.json"
LEVELS = [
    ("basic", "基础巩固", "⭐"),
    ("improve", "能力提高", "⭐⭐"),
    ("challenge", "拓展挑战", "⭐⭐⭐"),
]
EXAMPLE_TOPICS = {}


class PaperGenerationError(Exception):
    """面向用户的试卷生成错误。"""


class FriendlyArgumentParser(argparse.ArgumentParser):
    def error(self, message):
        print(f"参数错误：{message}", file=sys.stderr)
        print("请检查：学科、知识点、学生、年级、输出路径是否都已填写。", file=sys.stderr)
        print("示例：python3 scripts/generate_paper.py --subject 数学 --topic 最大公因数 --student 小明 --grade 五年级 --output 练习.md", file=sys.stderr)
        raise SystemExit(2)


def normalize_text(value, default):
    value = (value or "").strip()
    return value if value else default


def detect_stage(grade):
    if any(word in grade for word in ["初", "七", "八", "九"]):
        return "初中"
    if any(word in grade for word in ["高", "十", "十一", "十二"]):
        return "高中"
    if any(word in grade for word in ["一", "二", "三", "四", "五", "六", "小学"]):
        return "小学"
    return "通用"


def q(qtype, content, answer="", hint=""):
    item = {"type": qtype, "content": content}
    if answer:
        item["answer"] = answer
    if hint:
        item["hint"] = hint
    return item


def topic_has(topic, *keywords):
    return any(keyword.lower() in topic.lower() for keyword in keywords)


def load_balanced_bank():
    try:
        return json.loads(BALANCED_BANK_PATH.read_text(encoding="utf-8"))
    except FileNotFoundError:
        return {}
    except json.JSONDecodeError as exc:
        raise PaperGenerationError(f"均衡题库文件格式有误：{BALANCED_BANK_PATH}。请检查JSON格式。") from exc


BANK = load_balanced_bank()
EXAMPLE_TOPICS.update({
    subject: [item["topic"] for item in BANK.get(subject, [])]
    for subject in SUPPORTED_SUBJECTS
})


def trim_questions(questions, basic_count, improve_count, challenge_count):
    limits = {
        "basic": max(0, basic_count),
        "improve": max(0, improve_count),
        "challenge": max(0, challenge_count),
    }
    return {level: questions.get(level, [])[:limit] for level, limit in limits.items()}


def find_topic_profile(subject, topic):
    topic_lower = topic.lower()
    for profile in BANK.get(subject, []):
        candidates = [profile.get("topic", "")] + profile.get("aliases", [])
        if any(candidate and candidate.lower() in topic_lower for candidate in candidates):
            return profile
    return None


def generate_profile_questions(subject, topic, profile, basic_count=3, improve_count=3, challenge_count=2):
    key_points = profile.get("key_points", [])
    mistakes = profile.get("common_mistakes", [])
    focus = profile.get("question_focus", "基础理解、方法应用、综合迁移")
    canonical_topic = profile.get("topic", topic)
    stage = profile.get("stage", "小学/初中/高中")

    key_text = "、".join(key_points[:4]) if key_points else canonical_topic
    mistake_text = mistakes[0] if mistakes else "只记结论，没有说明理由"
    mistake_text_2 = mistakes[1] if len(mistakes) > 1 else mistake_text

    if subject in ["语文", "英语"]:
        questions = {
            "basic": [
                q("核心概念", f"围绕“{canonical_topic}”，写出本题最需要掌握的3个要点：{key_text}。"),
                q("基础练习", f"完成一道“{canonical_topic}”基础题，并在答案旁标出依据或关键词。"),
                q("易错辨析", f"判断并说明：学习“{canonical_topic}”时，最容易出现“{mistake_text}”这个问题。"),
            ],
            "improve": [
                q("方法应用", f"按“审题 -> 定位/分析 -> 作答”的顺序完成一道“{canonical_topic}”提高题。", hint=f"题型方向：{focus}。"),
                q("错题修正", f"给出一个“{canonical_topic}”错误答案，指出错因并改成规范答案。", hint=f"常见错因：{mistake_text_2}。"),
                q("表达训练", f"用一段完整的话总结“{canonical_topic}”的答题方法，要求有步骤、有例子。"),
            ],
            "challenge": [
                q("综合迁移", f"把“{canonical_topic}”放到一段新材料或新语境中，完成一道综合题并写出评分要点。"),
                q("自主命题", f"自己设计一道“{canonical_topic}”题目，给出参考答案和两个扣分点。"),
            ],
        }
    elif subject in ["数学", "物理", "化学"]:
        questions = {
            "basic": [
                q("概念填空", f"写出“{canonical_topic}”的核心概念/公式，并说明每个量的含义。", hint=f"重点：{key_text}。"),
                q("基础计算", f"完成一道“{canonical_topic}”基础计算题，要求写出公式、代入、结果和单位。"),
                q("判断辨析", f"判断并说明：出现“{mistake_text}”时，解题结果通常不可靠。"),
            ],
            "improve": [
                q("应用题", f"设计一个生活或实验情境，运用“{canonical_topic}”解决问题，并写出完整步骤。"),
                q("错因分析", f"给出一道“{canonical_topic}”错题，指出错在概念、公式、单位还是条件提取。", hint=f"常见错因：{mistake_text_2}。"),
                q("方法比较", f"尝试用两种方法解决同一道“{canonical_topic}”题，并说明哪种更适合考试。"),
            ],
            "challenge": [
                q("综合题", f"完成一道同时考查“{canonical_topic}”和另一个相关知识点的综合题。", hint=f"适用学段：{stage}。"),
                q("实验/探究", f"围绕“{canonical_topic}”设计一个验证或探究方案，写出变量、步骤和结论。"),
            ],
        }
    elif subject == "生物":
        questions = {
            "basic": [
                q("结构/概念", f"解释“{canonical_topic}”中的核心结构或概念，并说明功能。", hint=f"重点：{key_text}。"),
                q("功能匹配", f"把“{canonical_topic}”相关结构与功能一一对应，至少写出3组。"),
                q("易错判断", f"判断并说明：出现“{mistake_text}”时，说明概念还没有真正分清。"),
            ],
            "improve": [
                q("比较表", f"用表格比较“{canonical_topic}”中两个容易混淆的概念或结构。"),
                q("过程排序", f"把“{canonical_topic}”相关生命活动按正确顺序排列，并说明关键变化。"),
                q("实验分析", f"设计或分析一个观察“{canonical_topic}”的实验，写出材料、步骤、现象和结论。"),
            ],
            "challenge": [
                q("结构与功能", f"从“结构适应功能”的角度解释“{canonical_topic}”中的一个现象。"),
                q("综合探究", f"围绕“{canonical_topic}”提出一个可探究问题，写出变量、对照组和预期结果。"),
            ],
        }
    elif subject == "历史":
        questions = {
            "basic": [
                q("时间线索", f"整理“{canonical_topic}”的时间、人物/主体和关键事件。", hint=f"重点：{key_text}。"),
                q("概念解释", f"解释“{canonical_topic}”中的一个核心概念，并用一句史实支撑。"),
                q("易错判断", f"判断并说明：学习“{canonical_topic}”时，容易出现“{mistake_text}”。"),
            ],
            "improve": [
                q("原因影响", f"从背景、过程、影响中任选两个角度分析“{canonical_topic}”。"),
                q("史料分析", f"阅读一段关于“{canonical_topic}”的史料，提取观点并找出依据。"),
                q("比较归纳", f"把“{canonical_topic}”与相近历史事件比较，列出相同点和不同点。"),
            ],
            "challenge": [
                q("历史评价", f"评价“{canonical_topic}”的历史作用，要求兼顾进步性和局限性。"),
                q("时序迁移", f"把“{canonical_topic}”放入更长历史阶段中，说明它承前启后的意义。"),
            ],
        }
    elif subject == "地理":
        questions = {
            "basic": [
                q("概念识别", f"解释“{canonical_topic}”中的核心地理概念，并标出关键词。", hint=f"重点：{key_text}。"),
                q("图表判读", f"根据地图、坐标、气温降水图或统计图，提取“{canonical_topic}”相关信息。"),
                q("易错判断", f"判断并说明：出现“{mistake_text}”时，读图或分析会偏离结论。"),
            ],
            "improve": [
                q("空间定位", f"结合方位、区域或坐标信息，说明“{canonical_topic}”在地图上的判断步骤。"),
                q("区域比较", f"比较两个区域在“{canonical_topic}”上的差异，并写出自然或人文原因。"),
                q("图文转换", f"把一段关于“{canonical_topic}”的文字信息转成表格、示意图或步骤清单。"),
            ],
            "challenge": [
                q("综合分析", f"结合区域背景，分析“{canonical_topic}”对生产、生活或环境的影响。"),
                q("迁移应用", f"换一个陌生区域，说明如何用同样方法分析“{canonical_topic}”。"),
            ],
        }
    else:
        questions = {
            "basic": [
                q("概念解释", f"解释“{canonical_topic}”中的核心观点，并列出2个关键词。", hint=f"重点：{key_text}。"),
                q("情境判断", f"阅读一个生活情境，判断其中行为是否符合“{canonical_topic}”要求，并说明理由。"),
                q("易错辨析", f"判断并说明：出现“{mistake_text}”时，说明观点理解还不完整。"),
            ],
            "improve": [
                q("材料分析", f"结合一段社会生活材料，回答“{canonical_topic}”体现了什么观点，并找出材料依据。"),
                q("做法建议", f"遇到与“{canonical_topic}”相关的问题时，给出3条合法、具体、可执行的做法。"),
                q("观点辨析", f"对一个关于“{canonical_topic}”的片面观点进行辨析，要求先判断再说明理由。"),
            ],
            "challenge": [
                q("开放论述", f"围绕“{canonical_topic}”写一段150字左右的小论述，要求观点明确、联系生活、结论完整。"),
                q("行动方案", f"设计一个践行“{canonical_topic}”的小方案，写出目标、步骤和评价标准。"),
            ],
        }

    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_math_questions(topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    profile = find_topic_profile("数学", topic)
    if profile and not ("最大公因数" in topic or "公因数" in topic):
        return generate_profile_questions("数学", topic, profile, basic_count, improve_count, challenge_count)

    questions = {"basic": [], "improve": [], "challenge": []}

    if "最大公因数" in topic or "公因数" in topic:
        questions["basic"] = [
            q("计算", "用短除法求下列各组数的最大公因数。\n（1）18 和 27　（2）24 和 36　（3）45 和 60", "9；12；15"),
            q("填空", "如果 a 是 b 的倍数，那么 a 和 b 的最大公因数是（　　），最小公倍数是（　　）。", "b；a"),
            q("判断", "判断对错。\n（1）两个数的最大公因数一定比这两个数都小。（　　）\n（2）如果两个数互质，它们的最大公因数是1。（　　）", "×；√"),
        ]
        questions["improve"] = [
            q("应用题", "有两根铁丝，一根长 72 厘米，另一根长 54 厘米。要截成同样长的小段且没有剩余，每小段最长是多少厘米？一共可以截成多少段？", "18厘米；7段", "同样长、没有剩余、最长 -> 最大公因数。"),
            q("应用题", "48 本故事书和 64 本科技书平均分给若干个班，每班两类书数量都相同。最多可以分给多少个班？每班各多少本？", "16个班；故事书3本、科技书4本"),
            q("应用题", "长 84 厘米、宽 60 厘米的长方形彩纸剪成同样大的正方形且没有剩余，正方形边长最大是多少？最少剪成多少个？", "12厘米；35个"),
        ]
        questions["challenge"] = [
            q("思维题", "已知 a 和 b 的最大公因数是 12，最小公倍数是 72。如果 a = 36，那么 b = ？", "24"),
            q("思维题", "一个数除 60 余 4，除 80 也余 4。这个数最大是多少？", "28", "这个数能同时整除 60-4 和 80-4。"),
        ]
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    questions["basic"] = [
        q("概念填空", f"用自己的话解释“{topic}”的含义，并写出一个简单例子。"),
        q("基础计算", f"围绕“{topic}”设计并完成一道基础计算题，写出关键步骤。"),
        q("易错辨析", f"判断：学习“{topic}”时，只要记住公式就一定能做对所有题。（　　）", "×"),
    ]
    questions["improve"] = [
        q("应用题", f"把“{topic}”放进一个购物、路程或图形场景中，列式并解答。"),
        q("方法比较", f"同一道“{topic}”题目尝试用两种方法解决，并说明哪种更简洁。"),
        q("错因分析", f"写出一道“{topic}”相关错题，指出错误原因并改正。"),
    ]
    questions["challenge"] = [
        q("综合题", f"设计一道包含两个条件转换的“{topic}”综合题，并完整解答。"),
        q("开放题", f"给出一个生活中的“{topic}”问题，说明数学模型、解题步骤和答案。"),
    ]
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_chinese_questions(topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    profile = find_topic_profile("语文", topic)
    if profile:
        return generate_profile_questions("语文", topic, profile, basic_count, improve_count, challenge_count)

    if topic_has(topic, "阅读", "现代文"):
        questions = {
            "basic": [
                q("信息提取", "阅读短文后，用一句话概括主要内容，并圈出文中最能支持答案的一句话。", "评分要点：人物/对象、事件、结果要完整。"),
                q("词语理解", "解释文中一个关键词在语境中的意思，不能只写字典义。", "评分要点：联系上下文解释。"),
                q("中心判断", "判断文章主要表达的情感或观点，并写出两处依据。", "评分要点：观点明确，依据来自原文。"),
            ],
            "improve": [
                q("句子赏析", "从修辞、动词或描写角度赏析一个句子，按“手法+内容+效果”作答。", "示例框架：这句话运用了……，写出了……，表现了……。"),
                q("人物分析", "结合两处细节分析人物特点，不能只写形容词。", "评分要点：特点+细节+解释。"),
                q("标题作用", "分析标题的作用，至少从内容、线索、情感中选择两点回答。"),
            ],
            "challenge": [
                q("开放表达", "联系文章主题和生活实际，写一段80字左右的阅读感悟。", "评分要点：扣主题、有例子、有观点。"),
                q("迁移仿写", "仿照文中一个有表现力的句式写两句话，主题保持一致。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    if topic_has(topic, "病句"):
        questions = {
            "basic": [
                q("成分残缺", "修改病句：通过这次阅读活动，使我明白了坚持的重要。", "删去“通过”或“使”。"),
                q("搭配不当", "修改病句：同学们的学习态度和成绩都有了很大提高。", "“态度”应改为“端正”，或拆成两句。"),
                q("语序不当", "修改病句：我们认真讨论并听取了老师的建议。", "听取并讨论。"),
            ],
            "improve": [
                q("病因判断", "指出病句类型并修改：能否认真审题，是考试取得好成绩的关键。", "两面对一面；删去“能否”。"),
                q("综合修改", "找出这段话中的两处语病，并说明修改理由。"),
                q("表达优化", "把一个口语化句子改成书面表达，保持原意不变。"),
            ],
            "challenge": [
                q("段落修改", "修改一段含有语序、搭配、重复三类问题的小短文，并标注修改点。"),
                q("命题训练", "自己编写3个病句，分别对应成分残缺、搭配不当、语序不当，并给出答案。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    questions = {
        "basic": [
            q("字词积累", f"围绕“{topic}”整理 5 个关键词，分别写出意思或近义词。"),
            q("句子理解", f"用“{topic}”相关内容造 2 个完整句子，要求表达清楚。"),
            q("基础阅读", f"阅读一段与“{topic}”相关的短文，找出中心句并说明理由。"),
        ],
        "improve": [
            q("阅读分析", f"给一段“{topic}”相关材料，概括主要内容，并找出一个细节依据。"),
            q("表达训练", f"围绕“{topic}”写一段 120 字左右的小短文，要求有开头、经过、结尾。"),
            q("错题修正", f"修改一个与“{topic}”有关的病句，并说明病因。"),
        ],
        "challenge": [
            q("综合赏析", f"选择一段与“{topic}”相关的文字，从词语、修辞或情感中任选两点赏析。"),
            q("迁移写作", f"以“{topic}”为核心，写一个提纲：标题、中心、材料、结尾各一项。"),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_english_questions(topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    profile = find_topic_profile("英语", topic)
    if profile:
        return generate_profile_questions("英语", topic, profile, basic_count, improve_count, challenge_count)

    if topic_has(topic, "现在完成时", "present perfect"):
        questions = {
            "basic": [
                q("Choose", "I ____ my homework already.\nA. finish  B. finished  C. have finished  D. am finishing", "C", "already 常和现在完成时连用。"),
                q("Fill in", "She has ____ (live) in Beijing for three years.", "lived"),
                q("Transform", "He came here in 2022. 改写为现在完成时：He ____ ____ here since 2022.", "has been"),
            ],
            "improve": [
                q("Correction", "改错：I have bought this bike for two years.", "bought -> had 或改为 I bought this bike two years ago.", "延续时间要用延续性动词。"),
                q("Since/For", "用 since 或 for 填空：\n（1）____ last Monday  （2）____ two weeks  （3）____ 2020", "since；for；since"),
                q("Sentence", "用 have/has done 写3句关于自己学习经历的句子，每句包含 already、yet 或 ever 中的一个。"),
            ],
            "challenge": [
                q("Writing", "写一段60词左右的小短文，介绍你最近完成的一件事，至少使用3处现在完成时。"),
                q("Compare", "比较 I went to Shanghai. 和 I have been to Shanghai. 的意思差别，并各造一句。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    if topic_has(topic, "一般过去时", "past tense"):
        questions = {
            "basic": [
                q("Fill in", "Yesterday I ____ (visit) my grandparents.", "visited"),
                q("Choose", "She ____ to school by bike last week.\nA. goes  B. went  C. has gone", "B"),
                q("Negative", "Tom watched TV last night. 改为否定句。", "Tom didn't watch TV last night."),
            ],
            "improve": [
                q("Question", "They played football after school. 改为一般疑问句并作肯定回答。", "Did they play football after school? Yes, they did."),
                q("Timeline", "根据 yesterday、last Sunday、two days ago 各写一个一般过去时句子。"),
                q("Correction", "改错：Did you went to the park?", "went -> go"),
            ],
            "challenge": [
                q("Writing", "写一段80词左右日记，描述上周末做的三件事。"),
                q("Compare", "比较一般过去时和现在完成时的区别，并各举一个例句。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    questions = {
        "basic": [
            q("Vocabulary", f"Write 5 words or phrases about '{topic}' and give one Chinese meaning for each."),
            q("Sentence", f"Make 3 simple sentences about '{topic}'."),
            q("Choice", f"Choose the correct word to complete a sentence about '{topic}', then explain why."),
        ],
        "improve": [
            q("Grammar", f"Write 3 sentences about '{topic}' using the target grammar point, then mark the key structure."),
            q("Reading", f"Read a short paragraph about '{topic}' and answer: Who/What/Why?"),
            q("Correction", f"Find and correct 3 common mistakes in sentences about '{topic}'."),
        ],
        "challenge": [
            q("Writing", f"Write a 60-100 word paragraph about '{topic}' with at least 3 linking words."),
            q("Speaking", f"Prepare a 1-minute oral answer about '{topic}', including opinion and reason."),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_science_questions(subject, topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    profile = find_topic_profile(subject, topic)
    if profile:
        return generate_profile_questions(subject, topic, profile, basic_count, improve_count, challenge_count)

    if subject == "物理" and topic_has(topic, "浮力"):
        questions = {
            "basic": [
                q("概念填空", "浸在液体中的物体受到液体对它向上的力，这个力叫（　　）。", "浮力"),
                q("公式应用", "阿基米德原理公式可写为 F浮 = （　　）。", "G排 或 ρ液gV排"),
                q("判断", "物体漂浮时，浮力一定等于重力。（　　）", "√"),
            ],
            "improve": [
                q("计算", "一个物体排开水的体积为 2×10^-4 m³，水的密度为 1.0×10³ kg/m³，g取10 N/kg，求浮力。", "2 N"),
                q("现象分析", "为什么轮船是钢铁做的却能浮在水面上？", "关键：做成空心后排开水的体积增大，浮力可等于重力。"),
                q("实验探究", "用弹簧测力计测物体在空气中和水中的示数，如何求浮力？", "空气中示数 - 水中示数。"),
            ],
            "challenge": [
                q("综合分析", "同一木块分别漂浮在清水和盐水中，哪种液体中露出水面的体积更大？说明理由。", "盐水中更大；密度更大，需要排开更小体积即可平衡重力。"),
                q("变量控制", "设计实验探究浮力大小是否与液体密度有关，写出控制变量和改变变量。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    if subject == "化学" and topic_has(topic, "化学方程式", "方程式"):
        questions = {
            "basic": [
                q("配平", "配平：__H2 + __O2 = __H2O", "2H2 + O2 = 2H2O"),
                q("判断", "化学方程式配平的依据是质量守恒定律。（　　）", "√"),
                q("意义说明", "说出化学方程式 2H2 + O2 = 2H2O 表示的两层含义。", "反应物/生成物；粒子个数比或物质的量关系。"),
            ],
            "improve": [
                q("配平", "配平：__Fe + __O2 = __Fe3O4", "3Fe + 2O2 = Fe3O4"),
                q("纠错", "判断并修改：H2 + O2 = H2O2 可以表示氢气燃烧生成水。", "错误；生成水应为 2H2 + O2 = 2H2O。"),
                q("步骤说明", "写出配平化学方程式的三步：写反应式、配平、检查。"),
            ],
            "challenge": [
                q("综合配平", "配平并说明依据：__KMnO4 = __K2MnO4 + __MnO2 + __O2", "2KMnO4 = K2MnO4 + MnO2 + O2"),
                q("质量关系", "根据 2H2 + O2 = 2H2O，说明4份质量氢气完全反应需要多少份质量氧气。", "32份质量氧气"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    if subject == "生物" and topic_has(topic, "细胞"):
        questions = {
            "basic": [
                q("结构识别", "写出植物细胞特有的三个结构。", "细胞壁、液泡、叶绿体"),
                q("功能匹配", "细胞核的主要作用是什么？", "控制生命活动，储存遗传信息。"),
                q("判断", "动物细胞和植物细胞都有细胞膜、细胞质、细胞核。（　　）", "√"),
            ],
            "improve": [
                q("比较表", "用表格比较动物细胞和植物细胞的相同点与不同点。"),
                q("显微观察", "制作临时装片时为什么要滴清水或生理盐水？", "保持细胞正常形态。"),
                q("易错辨析", "为什么说细胞膜不是一堵完全封闭的墙？", "它能控制物质进出。"),
            ],
            "challenge": [
                q("结构与功能", "从结构适应功能角度说明叶绿体、线粒体分别为什么重要。"),
                q("实验设计", "设计一个观察洋葱表皮细胞的实验流程，并写出注意事项。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    unit = {
        "物理": ("概念/公式", "实验或现象", "单位和条件"),
        "化学": ("概念/方程式", "实验现象", "反应条件和守恒"),
        "生物": ("概念/结构", "生命现象", "条件和变量"),
    }[subject]
    questions = {
        "basic": [
            q("概念解释", f"解释“{topic}”中的一个核心{unit[0]}，并举一个例子。"),
            q("基础判断", f"判断一个关于“{topic}”的说法是否正确，并写出依据。"),
            q("图表整理", f"画出或列出“{topic}”的关键结构/步骤/关系图。"),
        ],
        "improve": [
            q("现象分析", f"描述一个与“{topic}”有关的{unit[1]}，说明原因。"),
            q("实验设计", f"设计一个验证“{topic}”的小实验，写出器材、步骤和观察点。"),
            q("易错辨析", f"列出学习“{topic}”时最容易混淆的两个概念，并比较不同点。"),
        ],
        "challenge": [
            q("综合应用", f"给出一个真实情境，运用“{topic}”解释结果，并标明{unit[2]}。"),
            q("探究题", f"围绕“{topic}”提出一个可探究问题，设计变量控制方案。"),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_humanities_questions(subject, topic, grade, basic_count=3, improve_count=3, challenge_count=2):
    profile = find_topic_profile(subject, topic)
    if profile:
        return generate_profile_questions(subject, topic, profile, basic_count, improve_count, challenge_count)

    if subject == "历史" and topic_has(topic, "辛亥革命"):
        questions = {
            "basic": [
                q("时间人物", "写出辛亥革命爆发的年份和重要领导人物。", "1911年；孙中山等。"),
                q("概念理解", "辛亥革命推翻了什么制度？", "清王朝统治和君主专制制度。"),
                q("判断", "辛亥革命使民主共和观念深入人心。（　　）", "√"),
            ],
            "improve": [
                q("原因分析", "从民族危机、清政府统治、革命思想传播中任选两点说明辛亥革命原因。"),
                q("影响分析", "说明辛亥革命的积极影响和局限性各一点。"),
                q("史料题", "阅读一段关于《中华民国临时约法》的材料，概括其体现的政治理念。", "民主共和、主权在民、限制权力等。"),
            ],
            "challenge": [
                q("评价题", "如何理解辛亥革命“既成功又失败”？请分两方面回答。"),
                q("比较题", "比较洋务运动、戊戌变法、辛亥革命在救国方式上的不同。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    if subject == "地理" and topic_has(topic, "气候"):
        questions = {
            "basic": [
                q("要素填空", "气候主要从（　　）和（　　）两个方面描述。", "气温；降水"),
                q("读图方法", "判断气候类型时，先看气温判断温度带，再看降水判断干湿季节。这个方法是否正确？", "正确"),
                q("概念辨析", "天气和气候有什么区别？", "天气时间短、变化快；气候时间长、较稳定。"),
            ],
            "improve": [
                q("类型判断", "某地全年高温多雨，最可能是什么气候类型？", "热带雨林气候"),
                q("成因分析", "为什么我国东部季风区夏季降水较多？", "受夏季风影响，来自海洋的湿润气流带来降水。"),
                q("图表题", "给一张气温曲线和降水柱状图，写出判读三步。", "看最冷月气温；看降水总量；看降水季节分配。"),
            ],
            "challenge": [
                q("区域比较", "比较温带季风气候和温带海洋性气候的气温、降水差异。"),
                q("生活联系", "说明气候对农业生产或民居形态的一种影响。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    if subject == "政治" and topic_has(topic, "公民权利", "权利"):
        questions = {
            "basic": [
                q("概念填空", "公民行使权利时，不得损害国家的、社会的、集体的利益和其他公民的合法权利。（　　）", "√"),
                q("列举", "列举两项公民基本权利。", "选举权和被选举权、受教育权、劳动权、人身自由等。"),
                q("判断", "权利和义务是完全分开的，享有权利就不用履行义务。（　　）", "×"),
            ],
            "improve": [
                q("情境分析", "同学在网络上随意发布他人照片并嘲笑，对方可以怎样维权？", "要求删除、赔礼道歉；必要时寻求老师、家长或法律帮助。"),
                q("做法题", "当自己的受教育权受到侵犯时，可以采取哪些合理方式？"),
                q("辨析题", "有人说“言论自由就是想说什么就说什么”。请辨析。", "错误；自由有边界，不能侵犯他人和公共利益。"),
            ],
            "challenge": [
                q("材料分析", "结合一个校园或网络生活案例，说明如何依法行使权利。"),
                q("观点表达", "围绕“权利与义务相统一”写一段120字左右的小论述。"),
            ],
        }
        return trim_questions(questions, basic_count, improve_count, challenge_count)

    focus = {
        "历史": ("时间、人物、事件", "原因和影响", "史料"),
        "地理": ("位置、气候、地形", "成因和影响", "地图或数据"),
        "政治": ("概念、观点、材料", "原因和做法", "时事材料"),
    }[subject]
    questions = {
        "basic": [
            q("基础梳理", f"整理“{topic}”中的 3 个关键词，并解释含义。"),
            q("信息提取", f"从一段关于“{topic}”的材料中提取{focus[0]}。"),
            q("判断说明", f"判断一个关于“{topic}”的说法是否正确，并写出理由。"),
        ],
        "improve": [
            q("材料分析", f"阅读一则关于“{topic}”的材料，概括核心观点并找出依据。"),
            q("因果分析", f"分析“{topic}”的{focus[1]}，至少写出两点。"),
            q("比较归纳", f"把“{topic}”与一个相近知识点作比较，列出相同点和不同点。"),
        ],
        "challenge": [
            q("综合探究", f"结合{focus[2]}，说明“{topic}”在现实或考试题中的应用。"),
            q("开放表达", f"围绕“{topic}”写一段 150 字左右的小论述，要求观点明确、依据充分。"),
        ],
    }
    return trim_questions(questions, basic_count, improve_count, challenge_count)


def generate_questions(subject, topic, grade, basic_count, improve_count, challenge_count):
    if subject == "数学":
        return generate_math_questions(topic, grade, basic_count, improve_count, challenge_count)
    if subject == "语文":
        return generate_chinese_questions(topic, grade, basic_count, improve_count, challenge_count)
    if subject == "英语":
        return generate_english_questions(topic, grade, basic_count, improve_count, challenge_count)
    if subject in ["物理", "化学", "生物"]:
        return generate_science_questions(subject, topic, grade, basic_count, improve_count, challenge_count)
    if subject in ["历史", "地理", "政治"]:
        return generate_humanities_questions(subject, topic, grade, basic_count, improve_count, challenge_count)
    raise PaperGenerationError(f"暂不支持“{subject}”。可用学科：{'、'.join(SUPPORTED_SUBJECTS)}。")


def validate_questions(questions):
    total = sum(len(items) for items in questions.values())
    if total == 0:
        raise PaperGenerationError("没有生成任何题目。请换一个更明确的知识点，例如“最大公因数”“现在完成时”“牛顿第二定律”。")


def build_output_data(args, questions):
    return {
        "subject": args.subject,
        "topic": args.topic,
        "student": args.student,
        "grade": args.grade,
        "stage": detect_stage(args.grade),
        "questions": questions,
        "usage_note": "九大学科均衡题库中的主题会优先生成三层梯度练习；其他主题使用学科兜底模板，正式给孩子使用前建议老师或家长快速浏览一遍。"
    }


def render_markdown(data):
    lines = [
        f"# {data['student']}的{data['subject']}练习",
        "",
        f"- 年级：{data['grade']}",
        f"- 知识点：{data['topic']}",
        f"- 学段：{data['stage']}",
        "",
        "> 先回顾讲解，再完成练习；遇到不会的题，先写出想到的步骤。",
        "",
    ]
    for key, title, star in LEVELS:
        lines.extend([f"## {title} {star}", ""])
        for index, item in enumerate(data["questions"].get(key, []), 1):
            lines.append(f"{index}. 【{item['type']}】{item['content']}")
            if item.get("hint"):
                lines.append(f"   - 提示：{item['hint']}")
            lines.append("")
    lines.extend(["## 参考答案", ""])
    for key, title, _ in LEVELS:
        answer_items = [item for item in data["questions"].get(key, []) if item.get("answer")]
        if not answer_items:
            continue
        lines.append(f"### {title}")
        for index, item in enumerate(answer_items, 1):
            lines.append(f"{index}. {item['answer']}")
        lines.append("")
    lines.append(f"> {data['usage_note']}")
    return "\n".join(lines)


def write_docx(path, data):
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise PaperGenerationError(
            "当前环境缺少 python-docx，无法生成 Word 文件。已建议改用 .md 输出，或先运行：python3 scripts/quick_setup.py"
        ) from exc

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    document.add_heading(f"{data['student']}的{data['subject']}练习", 0)
    document.add_paragraph(f"年级：{data['grade']}    学段：{data['stage']}    知识点：{data['topic']}")
    document.add_paragraph("先回顾讲解，再完成练习；遇到不会的题，先写出想到的步骤。")

    for key, title, star in LEVELS:
        document.add_heading(f"{title} {star}", level=1)
        for index, item in enumerate(data["questions"].get(key, []), 1):
            document.add_paragraph(f"{index}. 【{item['type']}】{item['content']}")
            if item.get("hint"):
                document.add_paragraph(f"提示：{item['hint']}")

    document.add_heading("参考答案", level=1)
    has_answers = False
    for key, title, _ in LEVELS:
        answer_items = [item for item in data["questions"].get(key, []) if item.get("answer")]
        if not answer_items:
            continue
        has_answers = True
        document.add_heading(title, level=2)
        for index, item in enumerate(answer_items, 1):
            document.add_paragraph(f"{index}. {item['answer']}")
    if not has_answers:
        document.add_paragraph("开放题暂无唯一答案，请根据讲解要点和表达完整度评价。")

    document.add_paragraph(data["usage_note"])
    document.save(path)


def write_output(output_path, data):
    path = Path(output_path).expanduser()
    path.parent.mkdir(parents=True, exist_ok=True)
    suffix = path.suffix.lower()

    if suffix == ".json":
        path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        return path
    if suffix == ".docx":
        try:
            write_docx(path, data)
            return path
        except PaperGenerationError:
            fallback = path.with_suffix(".md")
            fallback.write_text(render_markdown(data), encoding="utf-8")
            print(f"提示：Word 依赖不可用，已自动生成 Markdown 文件：{fallback}")
            return fallback

    if suffix not in [".md", ".txt"]:
        path = path.with_suffix(".md")
    path.write_text(render_markdown(data), encoding="utf-8")
    return path


def positive_int(value):
    try:
        number = int(value)
    except ValueError as exc:
        raise argparse.ArgumentTypeError("题目数量必须是整数。") from exc
    if number < 0 or number > 10:
        raise argparse.ArgumentTypeError("每层题目数量建议在 0-10 之间。")
    return number


def format_examples():
    lines = ["当前内置专项/推荐主题示例："]
    for subject in SUPPORTED_SUBJECTS:
        lines.append(f"- {subject}：{'、'.join(EXAMPLE_TOPICS[subject])}")
    lines.append("")
    lines.append("示例命令：")
    lines.append("python3 scripts/generate_paper.py --subject 物理 --topic 浮力 --student 小明 --grade 初二 --output 浮力练习.md")
    return "\n".join(lines)


def parse_args(argv):
    parser = FriendlyArgumentParser(
        description="生成K12三层练习题（支持九大学科，输出 docx/md/json）",
        epilog="示例：python3 scripts/generate_paper.py --subject 数学 --topic 最大公因数 --student 小明 --grade 五年级 --output 练习/小明.docx",
    )
    parser.add_argument("--list-examples", action="store_true", help="列出当前内置专项/推荐主题示例")
    parser.add_argument("--subject", help=f"学科：{'、'.join(SUPPORTED_SUBJECTS)}")
    parser.add_argument("--topic", help="知识点主题，例如：最大公因数、阅读理解、现在完成时")
    parser.add_argument("--student", help="学生姓名")
    parser.add_argument("--grade", help="年级，例如：小学五年级、初二、高一")
    parser.add_argument("--output", help="输出文件路径，支持 .docx/.md/.json")
    parser.add_argument("--basic-count", type=positive_int, default=3, help="基础巩固题数量，默认3")
    parser.add_argument("--improve-count", type=positive_int, default=3, help="能力提高题数量，默认3")
    parser.add_argument("--challenge-count", type=positive_int, default=2, help="拓展挑战题数量，默认2")
    args = parser.parse_args(argv)
    if args.list_examples:
        return args
    missing = [name for name in ["subject", "topic", "student", "grade", "output"] if not getattr(args, name)]
    if missing:
        raise PaperGenerationError(f"缺少必要参数：{', '.join('--' + name.replace('_', '-') for name in missing)}。可先运行 --list-examples 查看示例。")
    if args.subject not in SUPPORTED_SUBJECTS:
        raise PaperGenerationError(f"暂不支持“{args.subject}”。可用学科：{'、'.join(SUPPORTED_SUBJECTS)}。")
    return args


def main(argv=None):
    try:
        args = parse_args(argv or sys.argv[1:])
        if args.list_examples:
            print(format_examples())
            return 0
        args.topic = normalize_text(args.topic, "综合复习")
        args.student = normalize_text(args.student, "同学")
        args.grade = normalize_text(args.grade, "未提供年级")

        questions = generate_questions(
            args.subject,
            args.topic,
            args.grade,
            args.basic_count,
            args.improve_count,
            args.challenge_count,
        )
        validate_questions(questions)
        data = build_output_data(args, questions)
        saved_path = write_output(args.output, data)

        print(json.dumps(data, ensure_ascii=False, indent=2))
        print(f"\n练习已生成：{saved_path}")
        return 0
    except KeyboardInterrupt:
        print("\n已取消生成。", file=sys.stderr)
        return 130
    except PaperGenerationError as exc:
        print(f"生成失败：{exc}", file=sys.stderr)
        return 2
    except Exception as exc:
        print(f"生成失败：遇到未预期问题：{exc}", file=sys.stderr)
        print("建议：检查学科、知识点、输出路径是否正确；如仍失败，请把这条提示发给智能老师。", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
