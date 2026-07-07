#!/usr/bin/env python3
"""
PE Compliance Data Processing Script

This script provides data processing functions for the PECompliance skill,
including document parsing, deadline calculation, and risk assessment.

Requirements:
    pip install pdfplumber openpyxl python-docx pandas numpy
"""

import sys
import os
import re
import json
import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
import argparse

try:
    import pymupdf  # 极速PDF提取
    PY_MUPDF_SUPPORT = True
except ImportError:
    PY_MUPDF_SUPPORT = False
    print("Warning: pymupdf not installed. Fast PDF parsing disabled.")

try:
    import pdfplumber
    PDF_PLUMBER_SUPPORT = True
except ImportError:
    PDF_PLUMBER_SUPPORT = False
    print("Warning: pdfplumber not installed. PDF parsing disabled.")

try:
    import openpyxl
    EXCEL_SUPPORT = True
except ImportError:
    EXCEL_SUPPORT = False
    print("Warning: openpyxl not installed. Excel parsing disabled.")

try:
    import docx
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("Warning: python-docx not installed. Word parsing disabled.")

try:
    import pandas as pd
    PANDAS_SUPPORT = True
except ImportError:
    PANDAS_SUPPORT = False
    print("Warning: pandas not installed. Data analysis features limited.")


class PEComplianceProcessor:
    """Main processor for PE compliance data"""

    def __init__(self, jurisdiction: str = "US", fund_type: str = "Buyout"):
        self.jurisdiction = jurisdiction.upper()
        self.fund_type = fund_type
        self.regulatory_frameworks = self._load_frameworks()

    def _load_frameworks(self) -> Dict[str, Any]:
        """Load regulatory frameworks based on jurisdiction and fund type"""
        frameworks = {
            "US": {
                "SEC": {
                    "forms": ["Form ADV", "Form PF", "Form D"],
                    "deadlines": {
                        "Form ADV": {"annual": "03/31", "amendments": "as needed"},
                        "Form PF": {"quarterly": "end of quarter + 60 days"},
                    },
                    "requirements": [
                        "Registration with SEC or state",
                        "Advertising Rule compliance",
                        "Custody Rule compliance",
                        "Code of Ethics",
                        "Compliance Program (Rule 206(4)-7)"
                    ]
                },
                "AML": {
                    "requirements": [
                        "AML Program",
                        "Customer Identification Program (CIP)",
                        "Suspicious Activity Reporting (SAR)",
                        "OFAC screening"
                    ]
                },
                "TAX": {
                    "requirements": [
                        "FATCA reporting",
                        "K-1 distribution to investors",
                        "State tax filings"
                    ]
                }
            },
            "EU": {
                "AIFMD": {
                    "requirements": [
                        "AIFM authorization",
                        "Annex IV reporting",
                        "Depositary appointment",
                        "Transparency disclosures"
                    ],
                    "deadlines": {
                        "Annex IV": {"annual": "04/30", "semi-annual": "10/31"}
                    }
                },
                "AML": {
                    "requirements": [
                        "AML/CFT program",
                        "Customer Due Diligence (CDD)",
                        "Beneficial ownership identification"
                    ]
                },
                "ESG": {
                    "requirements": [
                        "SFDR disclosures (Principal Adverse Impacts)",
                        "Taxonomy alignment reporting",
                        "TCFD reporting"
                    ]
                }
            }
        }

        # Return frameworks for the selected jurisdiction
        return frameworks.get(self.jurisdiction, frameworks["US"])

    def parse_document(self, file_path: str) -> Dict[str, Any]:
        """Parse document and extract compliance-relevant information"""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        extracted_data = {
            "file_name": path.name,
            "file_type": path.suffix.lower(),
            "file_size": path.stat().st_size,
            "extracted_text": "",
            "compliance_keywords": [],
            "dates_found": [],
            "entities_found": []
        }

        # Extract text based on file type
        if path.suffix.lower() == '.pdf' and (PY_MUPDF_SUPPORT or PDF_PLUMBER_SUPPORT):
            extracted_data["extracted_text"] = self._extract_from_pdf(file_path)
        elif path.suffix.lower() == '.docx' and DOCX_SUPPORT:
            extracted_data["extracted_text"] = self._extract_from_docx(file_path)
        elif path.suffix.lower() in ['.xlsx', '.xls'] and EXCEL_SUPPORT:
            extracted_data["extracted_text"] = self._extract_from_excel(file_path)
        elif path.suffix.lower() == '.txt':
            extracted_data["extracted_text"] = self._extract_from_txt(file_path)
        else:
            print(f"Unsupported file type: {path.suffix}")
            return extracted_data

        # Analyze extracted text
        extracted_data["compliance_keywords"] = self._find_compliance_keywords(
            extracted_data["extracted_text"]
        )
        extracted_data["dates_found"] = self._find_dates(extracted_data["extracted_text"])
        extracted_data["entities_found"] = self._find_entities(extracted_data["extracted_text"])

        return extracted_data

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using pymupdf (fast) or pdfplumber"""
        text = ""

        # 优先使用 pymupdf (极速提取)
        if PY_MUPDF_SUPPORT:
            try:
                import pymupdf
                doc = pymupdf.open(file_path)
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    page_text = page.get_text()
                    if page_text:
                        text += page_text + "\n"
                doc.close()
                return text
            except Exception as e:
                print(f"Warning: pymupdf extraction failed, falling back to pdfplumber: {e}")

        # 回退到 pdfplumber
        if PDF_PLUMBER_SUPPORT:
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                print(f"Error extracting PDF with pdfplumber: {e}")

        return text

    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from Word document"""
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
        return text

    def _extract_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        text = ""
        try:
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            for sheet_name in wb.sheetnames:
                text += f"\n--- Sheet: {sheet_name} ---\n"
                ws = wb[sheet_name]
                for row in ws.iter_rows(values_only=True):
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    text += row_text + "\n"
        except Exception as e:
            print(f"Error extracting Excel: {e}")
        return text

    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            print(f"Error extracting text: {e}")
            return ""

    def _find_compliance_keywords(self, text: str) -> List[str]:
        """Find compliance-related keywords in text"""
        keywords = [
            "compliance", "regulation", "SEC", "AIFMD", "AML", "KYC", "FATCA",
            "CRS", "ESG", "risk", "governance", "audit", "reporting", "disclosure",
            "deadline", "filing", "form", "requirement", "obligation", "penalty"
        ]

        found = []
        text_lower = text.lower()
        for keyword in keywords:
            if keyword.lower() in text_lower:
                found.append(keyword)

        return found

    def _find_dates(self, text: str) -> List[str]:
        """Find dates in text (simple regex)"""
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # MM/DD/YYYY or DD/MM/YYYY
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY/MM/DD
            r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]* \d{1,2},? \d{4}'
        ]

        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            dates.extend(matches)

        return dates

    def _find_entities(self, text: str) -> List[str]:
        """Find potential entity names (simple heuristic)"""
        # Look for capitalized phrases
        entity_pattern = r'[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+'
        return re.findall(entity_pattern, text)

    def calculate_deadlines(self, start_date: str = None) -> Dict[str, str]:
        """Calculate compliance deadlines based on jurisdiction"""
        if start_date is None:
            start_date = datetime.date.today().isoformat()

        try:
            start = datetime.date.fromisoformat(start_date)
        except ValueError:
            start = datetime.date.today()

        deadlines = {}

        if self.jurisdiction == "US":
            # SEC Form ADV annual filing (March 31)
            adv_deadline = datetime.date(start.year, 3, 31)
            if start > adv_deadline:
                adv_deadline = datetime.date(start.year + 1, 3, 31)
            deadlines["Form ADV Annual"] = adv_deadline.isoformat()

            # Form PF quarterly (end of quarter + 60 days)
            # Simplified calculation
            quarter_month = ((start.month - 1) // 3) * 3 + 1
            quarter_end = datetime.date(start.year, quarter_month + 2, 30)
            pf_deadline = quarter_end + datetime.timedelta(days=60)
            deadlines["Form PF Quarterly"] = pf_deadline.isoformat()

        elif self.jurisdiction == "EU":
            # AIFMD Annex IV annual (April 30)
            annexiv_deadline = datetime.date(start.year, 4, 30)
            if start > annexiv_deadline:
                annexiv_deadline = datetime.date(start.year + 1, 4, 30)
            deadlines["AIFMD Annex IV Annual"] = annexiv_deadline.isoformat()

            # AIFMD Annex IV semi-annual (October 31)
            annexiv_semi = datetime.date(start.year, 10, 31)
            if start > annexiv_semi:
                annexiv_semi = datetime.date(start.year + 1, 10, 31)
            deadlines["AIFMD Annex IV Semi-Annual"] = annexiv_semi.isoformat()

        return deadlines

    def assess_risk(self, extracted_data: Dict[str, Any]) -> Dict[str, Any]:
        """Assess compliance risk based on extracted data"""
        risk_score = 0
        max_score = 100
        findings = []

        # Check for compliance keywords
        keywords_found = extracted_data.get("compliance_keywords", [])
        if len(keywords_found) < 3:
            risk_score += 20
            findings.append("Limited compliance terminology found in documents")

        # Check for dates
        dates_found = extracted_data.get("dates_found", [])
        if len(dates_found) == 0:
            risk_score += 15
            findings.append("No specific dates found for compliance deadlines")

        # Check document completeness
        text_length = len(extracted_data.get("extracted_text", ""))
        if text_length < 1000:
            risk_score += 10
            findings.append("Document appears brief; may lack detailed compliance information")

        # Calculate risk level
        risk_level = "Low"
        if risk_score > 50:
            risk_level = "High"
        elif risk_score > 25:
            risk_level = "Medium"

        return {
            "risk_score": risk_score,
            "risk_level": risk_level,
            "findings": findings,
            "recommendations": [
                "Review regulatory frameworks applicable to your fund",
                "Establish a compliance calendar with all deadlines",
                "Implement regular compliance training for staff",
                "Consider external compliance audit"
            ]
        }

    def generate_checklist(self) -> List[Dict[str, Any]]:
        """Generate compliance checklist based on jurisdiction and fund type"""
        checklist = []

        for framework, details in self.regulatory_frameworks.items():
            for requirement in details.get("requirements", []):
                checklist.append({
                    "framework": framework,
                    "requirement": requirement,
                    "status": "Not Started",
                    "deadline": "",
                    "responsible": "Compliance Officer",
                    "notes": ""
                })

        return checklist

    def build_compliance_prompt(self, extracted_text: str, file_name: str = "") -> str:
        """构建包含最新监管条例检索要求的提示词"""
        base_prompt = f"""你是一名拥有15年经验的顶尖私募股权（PE/VC）基金法务与合规总监。

## 审查文本
文件：{file_name}
内容预览：{extracted_text[:2000]}...

## 监管对齐要求（必须遵守）
你必须检索并对比最新的私募基金监管条例，包括但不限于：
1. 《私募投资基金监督管理条例》（国务院令第XXX号）
2. 中国证券投资基金业协会（AMAC）最新发布的：
   - 《私募投资基金登记备案办法》
   - 《私募投资基金募集行为管理办法》
   - 《私募投资基金信息披露管理办法》
   - 《私募基金管理人内部控制指引》
3. 其他相关自律规则和监管问答

## 审查模块（8大支柱）
基于以下8个模块进行穿透式审查：
1. 登记备案模块：高管资质、实际控制人、出资结构
2. 基金募集模块：公开推介、保本保收益承诺、合格投资者认定
3. 投资者适当性：双录要求、风险评级匹配
4. 管理人内部控制：利益冲突防范、防火墙、关联交易
5. 基金投资运作：明股实债、资金池、期限错配、杠杆限制
6. 信息披露模块：重大事项报告时效、定期报告准确性
7. 基金托管人：托管协议条款、资金划拨权限隔离
8. 基金合同/合伙协议：管理费、门槛收益、Carry分配、关键人士条款、违约责任

## 输出要求
必须生成简体中文Word报告（私募合规报告.docx），包含：
1. 【1】合规健康度概览：风险等级（高/中/低）+ 最致命"否决项"
2. 【2】红线违规预警：直接违反现行监管规定的条款
3. 【3】操作风险排查明细（基于8大模块）：每个模块的瑕疵与漏洞
4. 【4】合规整改与商业博弈建议：GP/LP博弈角度的条款修改建议

## 分析原则
- 穿透审查：代持、多层嵌套必须预警
- 商业与合规并重：提供替代性合规方案

现在，基于最新监管条例，对上述文本进行合规审查，并生成完整报告。"""

        return base_prompt

    def generate_word_report(self, analysis_results: Dict[str, Any], output_path: str = None) -> str:
        """生成带层级标题的Word文档报告"""
        if not DOCX_SUPPORT:
            print("Error: python-docx not installed. Cannot generate Word report.")
            return ""

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE

            # 创建文档
            doc = Document()

            # 添加标题
            title = doc.add_heading('私募基金合规审查报告', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加报告信息
            doc.add_paragraph(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph(f"审查文件：{analysis_results.get('file_name', 'N/A')}")
            doc.add_paragraph(f"文件类型：{analysis_results.get('file_type', 'N/A')}")
            doc.add_paragraph(f"文件大小：{analysis_results.get('file_size', 0)} bytes")
            doc.add_paragraph("")

            # 1. 合规健康度概览
            doc.add_heading('【1】合规健康度概览', level=1)
            risk_level = analysis_results.get('risk_assessment', {}).get('risk_level', '未知')
            risk_score = analysis_results.get('risk_assessment', {}).get('risk_score', 0)

            p = doc.add_paragraph()
            p.add_run(f"总体合规风险等级：").bold = True
            p.add_run(f" {risk_level}（风险分数：{risk_score}/100）")

            # 致命否决项
            doc.add_heading('最致命"否决项"', level=2)
            findings = analysis_results.get('risk_assessment', {}).get('findings', [])
            if findings:
                for finding in findings:
                    doc.add_paragraph(f"• {finding}", style='ListBullet')
            else:
                doc.add_paragraph("未发现致命否决项")

            # 2. 红线违规预警
            doc.add_heading('【2】红线违规预警（重大风险）', level=1)
            keywords = analysis_results.get('compliance_keywords', [])
            if keywords:
                doc.add_paragraph("检测到以下合规相关关键词：")
                for keyword in keywords:
                    doc.add_paragraph(f"• {keyword}", style='ListBullet')
            else:
                doc.add_paragraph("未检测到明显的合规关键词")

            # 日期发现
            dates = analysis_results.get('dates_found', [])
            if dates:
                doc.add_paragraph("文档中发现以下关键日期：")
                for date in dates[:10]:  # 最多显示10个日期
                    doc.add_paragraph(f"• {date}", style='ListBullet')

            # 3. 操作风险排查明细（8大模块）
            doc.add_heading('【3】操作风险排查明细（基于8大模块）', level=1)
            modules = [
                "1. 登记备案模块",
                "2. 基金募集模块",
                "3. 投资者适当性",
                "4. 管理人内部控制",
                "5. 基金投资运作",
                "6. 信息披露模块",
                "7. 基金托管人",
                "8. 基金合同/合伙协议"
            ]

            for module in modules:
                doc.add_heading(module, level=2)
                doc.add_paragraph("审查要点：需结合具体文档内容分析")
                doc.add_paragraph("风险等级：待评估")
                doc.add_paragraph("")

            # 4. 合规整改与商业博弈建议
            doc.add_heading('【4】合规整改与商业博弈建议', level=1)
            recommendations = analysis_results.get('risk_assessment', {}).get('recommendations', [])
            if recommendations:
                doc.add_paragraph("基于当前分析，提出以下建议：")
                for rec in recommendations:
                    doc.add_paragraph(f"• {rec}", style='ListBullet')
            else:
                doc.add_paragraph("请结合具体文档内容提供定制化建议")

            # 添加监管依据
            doc.add_heading('监管依据', level=2)
            regulations = [
                "《私募投资基金监督管理条例》",
                "《私募投资基金登记备案办法》（AMAC）",
                "《私募投资基金募集行为管理办法》（AMAC）",
                "《私募投资基金信息披露管理办法》（AMAC）",
                "《私募基金管理人内部控制指引》（AMAC）"
            ]
            for reg in regulations:
                doc.add_paragraph(f"• {reg}", style='ListBullet')

            # 保存文档
            if output_path is None:
                desktop = os.path.join(os.path.expanduser('~'), 'Desktop')
                output_path = os.path.join(desktop, '私募合规报告.docx')

            doc.save(output_path)
            return output_path

        except Exception as e:
            print(f"Error generating Word report: {e}")
            return ""


def main():
    """Command-line interface for the PE Compliance processor"""
    parser = argparse.ArgumentParser(description="PE Compliance Data Processor")
    parser.add_argument("--file", type=str, help="Path to document for analysis")
    parser.add_argument("--jurisdiction", type=str, default="US",
                       choices=["US", "EU", "UK", "CN", "Global"],
                       help="Jurisdiction for compliance analysis")
    parser.add_argument("--fund-type", type=str, default="Buyout",
                       choices=["Buyout", "Growth", "Venture", "Distressed", "RealEstate", "Infrastructure"],
                       help="Type of private equity fund")
    parser.add_argument("--checklist", action="store_true",
                       help="Generate compliance checklist")
    parser.add_argument("--deadlines", action="store_true",
                       help="Calculate compliance deadlines")
    parser.add_argument("--risk", action="store_true",
                       help="Assess compliance risk")
    parser.add_argument("--output", type=str, default="json",
                       choices=["json", "text", "csv"],
                       help="Output format")
    parser.add_argument("--word-report", action="store_true",
                       help="Generate Word report document")

    args = parser.parse_args()

    # Initialize processor
    processor = PEComplianceProcessor(
        jurisdiction=args.jurisdiction,
        fund_type=args.fund_type
    )

    results = {}

    # Process file if provided
    if args.file:
        print(f"Analyzing document: {args.file}")
        extracted_data = processor.parse_document(args.file)
        results["document_analysis"] = extracted_data

        if args.risk:
            risk_assessment = processor.assess_risk(extracted_data)
            results["risk_assessment"] = risk_assessment

        # 构建综合分析结果用于Word报告
        if args.word_report:
            # 确保有风险评估结果
            if args.risk:
                assessment = risk_assessment
            else:
                assessment = processor.assess_risk(extracted_data)

            combined_analysis = {
                **extracted_data,
                "risk_assessment": assessment,
                "jurisdiction": args.jurisdiction,
                "fund_type": args.fund_type
            }
            report_path = processor.generate_word_report(combined_analysis)
            if report_path:
                print(f"Word报告已生成：{report_path}")
                results["word_report_path"] = report_path

            # 构建合规审查提示词
            prompt = processor.build_compliance_prompt(
                extracted_data.get("extracted_text", ""),
                extracted_data.get("file_name", "")
            )
            results["compliance_prompt"] = prompt[:1000] + "..." if len(prompt) > 1000 else prompt

    # Generate checklist if requested
    if args.checklist:
        checklist = processor.generate_checklist()
        results["compliance_checklist"] = checklist

    # Calculate deadlines if requested
    if args.deadlines:
        deadlines = processor.calculate_deadlines()
        results["compliance_deadlines"] = deadlines

    # Output results
    if args.output == "json":
        print(json.dumps(results, indent=2, ensure_ascii=False))
    elif args.output == "text":
        for key, value in results.items():
            print(f"\n=== {key.upper().replace('_', ' ')} ===")
            if isinstance(value, list):
                for item in value:
                    print(f"  - {item}")
            elif isinstance(value, dict):
                for subkey, subvalue in value.items():
                    print(f"  {subkey}: {subvalue}")
            else:
                print(f"  {value}")
    elif args.output == "csv" and PANDAS_SUPPORT:
        df = pd.json_normalize(results)
        print(df.to_csv(index=False))

    return 0


if __name__ == "__main__":
    sys.exit(main())