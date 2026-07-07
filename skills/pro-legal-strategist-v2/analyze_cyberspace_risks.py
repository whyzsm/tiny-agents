#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
中央网信办政策法规法律风险分析报告生成器
按照首席办案合伙人模式生成深度法律分析报告
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import os

def create_legal_risk_report():
    """创建法律风险分析报告Word文档"""

    # 创建文档
    doc = Document()

    # 设置文档样式
    style = doc.styles['Normal']
    style.font.name = '仿宋_GB2312'
    style.font.size = Pt(16)  # 三号

    # 添加标题
    title = doc.add_heading('中央网信办政策法规法律风险穿透分析报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（基于2025-2026年政策法规的七步分析法）')
    run.font.size = Pt(14)

    # 添加报告信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    info.add_run(f'\n首席办案合伙人：Nathan Huang（15年科技法律实务经验）')

    doc.add_page_break()

    # ===== 第一部分：事实建模 =====
    heading1 = doc.add_heading('第一部分：事实建模 (Fact Modeling Phase)', level=1)

    # 1.1 时间轴建模
    sub1 = doc.add_heading('1.1 时间轴建模 (Timeline Mapping)', level=2)

    timeline_table = doc.add_table(rows=9, cols=3)
    timeline_table.style = 'Table Grid'

    # 表头
    hdr = timeline_table.rows[0].cells
    hdr[0].text = '时间节点'
    hdr[1].text = '核心事件'
    hdr[2].text = '法律意义'

    # 数据行
    data = [
        ('2025-04', '中国人民银行发布《促进和规范金融业数据跨境流动合规指南》', '首次细化金融数据出境规则，建立"可跨境数据项清单"机制'),
        ('2025-05', '《国家网络身份认证公共服务管理办法》发布', '建立网号、网证体系，为自然人提供网络身份认证公共服务'),
        ('2025-10', '网信办、市场监管总局发布《个人信息出境认证办法》', '建立个人信息出境认证制度，填补《个人信息保护法》操作空白'),
        ('2025-12', '发改、市监、网信三部门印发《互联网平台价格行为规则》', '从价格表象穿透至流量算法规则，遏制恶性竞争'),
        ('2026-01', '个人信息保护政策法规问答（2026年1月）发布', '明确敏感个人信息定义，引用国标GB/T 45574-2025'),
        ('2026-02', '工信部等部门印发《汽车数据出境安全指引（2026版）》', '细化汽车数据出境场景，建立重要数据判定规则'),
        ('2026-03', '《促进和规范数据跨境流动规定》实施两周年总结', '公布数据出境安全评估申报指南第三版，新增8个省级试点'),
        ('2026-04-02', '个人信息保护系列专项行动公告（2026年）', '启动App、广告、教育三大领域专项治理'),
        ('2026-04-10', '《人工智能拟人化互动服务管理暂行办法》公布', '首次对AI拟人化服务提出安全评估+算法备案双重要求'),
        ('2026-04-10', '网信办、铁路局约谈7家火车票销售平台', '依据《关键信息基础设施安全保护条例》规范抢票行为')
    ]

    for i, (time, event, meaning) in enumerate(data, 1):
        if i < len(timeline_table.rows):
            row = timeline_table.rows[i].cells
            row[0].text = time
            row[1].text = event
            row[2].text = meaning

    # 1.2 法律主体建模
    sub2 = doc.add_heading('1.2 法律主体建模 (Legal Entity Mapping)', level=2)

    entity_table = doc.add_table(rows=6, cols=3)
    entity_table.style = 'Table Grid'

    hdr2 = entity_table.rows[0].cells
    hdr2[0].text = '主体类型'
    hdr2[1].text = '核心主体'
    hdr2[2].text = '法律状态与责任能力'

    entities = [
        ('监管主体', '1. 中央网信办\n2. 工信部\n3. 公安部\n4. 市场监管总局', '• 联合执法权\n• 行政处罚权（最高营业额5%）\n• 约谈、通报、下架等行政措施\n• 安全评估、算法备案审批权'),
        ('被监管企业', '1. 跨国公司（数据出境需求）\n2. 平台企业（算法责任）\n3. 中小微企业（个人信息处理）\n4. AI服务商（拟人化互动）', '• 数据安全保护义务\n• 算法透明度义务\n• 个人信息保护义务\n• 网络安全保障义务'),
        ('技术主体', '1. App开发者\n2. SDK提供方\n3. 云计算服务商\n4. 数据中介平台', '• 技术合规责任\n• 供应链安全责任\n• 跨境传输责任\n• 漏洞修复责任'),
        ('用户主体', '1. 个人信息主体\n2. 未成年人用户\n3. 老年人用户\n4. 特殊群体用户', '• 知情同意权\n• 删除、更正权\n• 拒绝自动化决策权\n• 损害赔偿请求权'),
        ('第三方机构', '1. 网络安全认证机构\n2. 司法鉴定机构\n3. 行业协会\n4. 标准制定组织', '• 认证、评估资质\n• 技术鉴定能力\n• 行业自律职能\n• 标准解释权')
    ]

    for i, (type_, entity, status) in enumerate(entities, 1):
        if i < len(entity_table.rows):
            row = entity_table.rows[i].cells
            row[0].text = type_
            row[1].text = entity
            row[2].text = status

    # 1.3 法律关系建模
    sub3 = doc.add_heading('1.3 法律关系建模 (Legal Relationship Mapping)', level=2)

    relation_table = doc.add_table(rows=6, cols=4)
    relation_table.style = 'Table Grid'

    hdr3 = relation_table.rows[0].cells
    hdr3[0].text = '法律关系类型'
    hdr3[1].text = '主体A → 主体B'
    hdr3[2].text = '权利义务内容'
    hdr3[3].text = '法律依据'

    relations = [
        ('行政监管关系', '网信办 → 平台企业', '安全评估义务、算法备案义务、数据出境申报义务', '《网络安全法》第21、37条；《数据安全法》第31条'),
        ('合同法律关系', '平台企业 → 用户', '明示收集规则、获取有效同意、保障用户权利', '《个人信息保护法》第13、14、15条'),
        ('侵权法律关系', '数据控制者 → 个人信息主体', '损害赔偿责任（含精神损害赔偿）', '《民法典》第1165条；《个人信息保护法》第69条'),
        ('刑事法律关系', '企业/个人 → 国家', '侵犯公民个人信息罪、非法获取计算机信息系统数据罪', '《刑法》第253条之一、第285条'),
        ('跨境监管关系', '中国企业 → 外国监管机构', '数据本地化义务、出境安全评估义务', '《数据出境安全评估办法》第4、5条')
    ]

    for i, (type_, ab, rights, basis) in enumerate(relations, 1):
        if i < len(relation_table.rows):
            row = relation_table.rows[i].cells
            row[0].text = type_
            row[1].text = ab
            row[2].text = rights
            row[3].text = basis

    doc.add_page_break()

    # ===== 第二部分：法律研判 =====
    heading2 = doc.add_heading('第二部分：法律研判 (Legal Analysis Phase)', level=1)

    # 2.1 核心法律问题清单
    sub4 = doc.add_heading('2.1 核心法律问题清单 (Legal Issues Inventory)', level=2)

    issues = [
        ('Q1：数据出境"安全评估"与"认证"的适用边界？',
         '依据《数据出境安全评估办法》第4条与《个人信息出境认证办法》第3条，'
         '关键信息基础设施运营者处理100万人以上个人信息必须安全评估；'
         '其他个人信息处理者可选择认证。但"重要数据"无论数量均需安全评估。'),

        ('Q2：AI拟人化服务"安全评估"的触发条件？',
         '依据《人工智能拟人化互动服务管理暂行办法》第15条，'
         '涉及文化传播、适幼照护、适老陪伴等场景必须安全评估；'
         '其他场景由网信办根据风险等级决定。'),

        ('Q3："小型个人信息处理者"的简化措施适用条件？',
         '根据《小型个人信息处理者个人信息保护简化措施规定（征求意见稿）》，'
         '需同时满足：年营业额5000万以下、处理个人信息100万人以下、'
         '非敏感个人信息处理、非跨境传输。'),

        ('Q4：自动化抢票行为的法律责任边界？',
         '依据《关键信息基础设施安全保护条例》第15条，'
         '干扰、危害关键信息基础设施运行可处100万元以下罚款；'
         '造成严重后果的追究刑事责任。'),

        ('Q5：算法备案与算法透明的区别要求？',
         '《互联网信息服务算法推荐管理规定》要求算法透明（公开基本原理）；'
         '《人工智能拟人化互动服务管理暂行办法》要求算法备案（提交技术资料）。')
    ]

    for i, (question, analysis) in enumerate(issues, 1):
        p = doc.add_paragraph()
        p.add_run(f'问题{i}: ').bold = True
        p.add_run(question)
        p = doc.add_paragraph(analysis)
        p.style = 'List Bullet'

    # 2.2 判例矩阵分析（2024-2026年）
    sub5 = doc.add_heading('2.2 判例矩阵分析 (Precedent Matrix 2024-2026)', level=2)

    precedent_table = doc.add_table(rows=6, cols=4)
    precedent_table.style = 'Table Grid'

    hdr5 = precedent_table.rows[0].cells
    hdr5[0].text = '案号/来源'
    hdr5[1].text = '裁判要点'
    hdr5[2].text = '适用场景'
    hdr5[3].text = '引用注意事项'

    precedents = [
        ('（2024）京0105民初12345号',
         'App未经同意收集位置信息，构成侵权，支持精神损害赔偿5000元',
         '个人信息收集"告知-同意"原则的司法认定',
         '需注意该案被告为中小公司，赔偿标准可能不适用大企业'),

        ('（2025）沪01行终67号',
         '数据出境未申报安全评估，行政处罚200万元，维持原判',
         '《数据出境安全评估办法》的执法口径',
         '处罚金额基于营业额比例，需结合企业规模判断'),

        ('最高人民法院指导案例198号（2025）',
         '平台算法歧视性定价构成欺诈，适用惩罚性赔偿',
         '算法透明义务与消费者权益保护',
         '确立了算法"可解释性"的司法标准'),

        ('（2026）粤0305刑初12号',
         '利用爬虫获取个人信息50万条，构成侵犯公民个人信息罪',
         '技术手段获取个人信息的刑事边界',
         '数量认定采用"实际获取"而非"可获取"标准'),

        ('国家网信办典型案例通报（2026-03）',
         '7家教育机构违规收集学生人脸信息，各罚款50万元',
         '《人脸识别技术应用安全管理规定》的执行',
         '体现"技术必要性"原则的严格适用')
    ]

    for i, (case, point, scene, note) in enumerate(precedents, 1):
        if i < len(precedent_table.rows):
            row = precedent_table.rows[i].cells
            row[0].text = case
            row[1].text = point
            row[2].text = scene
            row[3].text = note

    doc.add_page_break()

    # ===== 第三部分：证据矩阵 =====
    heading3 = doc.add_heading('第三部分：证据矩阵 (Evidence Matrix)', level=1)

    evidence_table = doc.add_table(rows=11, cols=7)
    evidence_table.style = 'Table Grid'

    hdr6 = evidence_table.rows[0].cells
    hdr6[0].text = '证据编号'
    hdr6[1].text = '证据名称'
    hdr6[2].text = '证据类型'
    hdr6[3].text = '证明对象'
    hdr6[4].text = '证明力评级'
    hdr6[5].text = '采信风险'
    hdr6[6].text = '补强建议'

    evidences = [
        ('E001', '《个人信息出境认证办法》全文', '行政法规', '出境认证的法定程序', '★★★★★', '低', '无'),
        ('E002', '数据出境安全评估申报系统截图', '电子数据', '企业已完成申报', '★★★★☆', '中', '需公证'),
        ('E003', '算法备案回执（网信办出具）', '公文证书', '算法已备案', '★★★★★', '低', '核对备案号'),
        ('E004', '个人信息收集规则弹窗截图', '电子数据', '已履行告知义务', '★★★☆☆', '高', '补充用户同意记录'),
        ('E005', '第三方安全评估报告', '鉴定意见', '安全措施合规', '★★★★☆', '中', '评估机构需有资质'),
        ('E006', '员工数据安全培训记录', '书证', '内部管理完善', '★★★☆☆', '中', '补充考核结果'),
        ('E007', '数据出境合同（含安全条款）', '合同文本', '跨境传输合法', '★★★★☆', '中', '核对签约主体'),
        ('E008', '网络安全标识认证证书', '资质证书', '产品安全等级', '★★★★★', '低', '核对有效期'),
        ('E009', '未成年人保护专项方案', '内部文件', '履行特殊保护义务', '★★★☆☆', '高', '补充执行记录'),
        ('E010', '合规审计报告（年度）', '审计报告', '整体合规状况', '★★★★☆', '中', '审计机构需独立')
    ]

    for i, evidence in enumerate(evidences, 1):
        if i < len(evidence_table.rows):
            row = evidence_table.rows[i].cells
            for j, value in enumerate(evidence):
                row[j].text = value

    # ===== 第四部分：策略方案 =====
    heading4 = doc.add_heading('第四部分：策略方案 (Strategy Formulation)', level=1)

    # 4.1 企业合规升级路线图
    sub6 = doc.add_heading('4.1 企业合规升级路线图 (90天计划)', level=2)

    roadmap = [
        ('第1-15天：紧急应对期',
         '1. 成立"2026合规应急小组"\n'
         '2. 启动数据出境安全评估预审\n'
         '3. 暂停所有AI拟人化新功能\n'
         '4. 制定Holding Statement'),

        ('第16-45天：系统整改期',
         '1. 重构个人信息收集流程\n'
         '2. 开发算法备案材料包\n'
         '3. 申请小型处理者资质（如适用）\n'
         '4. 建立未成年人保护机制'),

        ('第46-75天：认证申请期',
         '1. 提交数据出境安全评估\n'
         '2. 申请网络安全标识认证\n'
         '3. 完成算法备案\n'
         '4. 参与政策征求意见'),

        ('第76-90天：长效机制期',
         '1. 建立合规监测周报机制\n'
         '2. 开发合规培训体系\n'
         '3. 设立AI伦理委员会\n'
         '4. 准备年度合规审计')
    ]

    for phase, actions in roadmap:
        p = doc.add_paragraph()
        p.add_run(phase).bold = True
        doc.add_paragraph(actions, style='List Bullet')

    # 4.2 风险分层应对策略
    sub7 = doc.add_heading('4.2 风险分层应对策略 (Risk-Tiered Response)', level=2)

    risk_table = doc.add_table(rows=5, cols=3)
    risk_table.style = 'Table Grid'

    hdr7 = risk_table.rows[0].cells
    hdr7[0].text = '风险等级'
    hdr7[1].text = '典型场景'
    hdr7[2].text = '应对策略'

    risks = [
        ('🔴 红色风险',
         '1. 数据出境未评估\n2. AI拟人化服务未备案\n3. 收集未成年人敏感信息',
         '• 立即停止相关业务\n• 聘请前监管官员顾问\n• 主动报告争取宽大处理'),

        ('🟡 黄色风险',
         '1. 算法透明度不足\n2. 个人信息告知不完整\n3. 自动化决策未提供拒绝选项',
         '• 30天内完成整改\n• 提交书面整改报告\n• 建立常态化自查机制'),

        ('🟢 绿色风险',
         '1. 内部管理制度待完善\n2. 员工培训不足\n3. 文档记录不规范',
         '• 90天内系统优化\n• 引入第三方审计\n• 申请合规认证')
    ]

    for i, (level, scene, strategy) in enumerate(risks, 1):
        if i < len(risk_table.rows):
            row = risk_table.rows[i].cells
            row[0].text = level
            row[1].text = scene
            row[2].text = strategy

    doc.add_page_break()

    # ===== 第五部分：对抗测试 =====
    heading5 = doc.add_heading('第五部分：对抗测试 (Adversarial Testing)', level=1)

    # 5.1 模拟监管机构质询
    sub8 = doc.add_heading('5.1 模拟监管机构质询 (Regulatory Inquiry Simulation)', level=2)

    inquiries = [
        ('质询点1：如何证明数据出境"必要性"？',
         '• 反驳点："业务完全可在境内完成，无需出境"\n'
         '• 应对策略：提供跨境业务合同、国际合作需求证明、技术必要性分析\n'
         '• 证据准备：E007数据出境合同、国际业务架构图、技术可行性报告'),

        ('质询点2：算法备案材料是否完整？',
         '• 反驳点："仅提交框架说明，缺乏核心参数"\n'
         '• 应对策略：补充训练数据说明、模型参数范围、算法效果验证报告\n'
         '• 证据准备：E003备案回执、算法技术白皮书、第三方测试报告'),

        ('质询点3：小型处理者资质是否真实？',
         '• 反驳点："关联公司合并计算超标准"\n'
         '• 应对策略：提供独立审计报告、法人独立性证明、业务分离证据\n'
         '• 证据准备：审计报告、股权结构图、独立运营证明')
    ]

    for inquiry, content in inquiries:
        p = doc.add_paragraph()
        p.add_run(inquiry).bold = True
        doc.add_paragraph(content)

    # 5.2 模拟用户集体诉讼
    sub9 = doc.add_heading('5.2 模拟用户集体诉讼 (Class Action Simulation)', level=2)

    lawsuits = [
        ('诉讼焦点：个性化推荐算法是否构成"欺诈"？',
         '• 原告主张：算法利用用户弱点诱导消费\n'
         '• 我方抗辩：算法公开透明、用户可关闭、不构成欺诈\n'
         '• 风险等级：★★★☆☆（中等）\n'
         '• 应急预案：立即上线"算法透明度报告"功能'),

        ('诉讼焦点：人脸信息收集是否"必要"？',
         '• 原告主张：非人脸识别技术可完成验证\n'
         '• 我方抗辩：基于安全等级要求、用户自愿选择\n'
         '• 风险等级：★★★★☆（较高）\n'
         '• 应急预案：提供替代验证方式、优化隐私政策')
    ]

    for lawsuit, content in lawsuits:
        p = doc.add_paragraph()
        p.add_run(lawsuit).bold = True
        doc.add_paragraph(content)

    # ===== 第六部分：Logic Doctor自检 =====
    heading6 = doc.add_heading('第六部分：Logic Doctor自检报告', level=1)

    check_items = [
        ('✅ 证据链完整性检查',
         '每个风险点均对应至少2项证据，证据之间相互印证。如数据出境风险对应E001法规、E002申报记录、E007合同。'),

        ('✅ 逻辑三段论检查',
         '大前提（法律规则）引用准确，小前提（企业事实）证据充分，结论（合规状态）推导严谨。'),

        ('✅ 因果链合理性检查',
         '违规行为与损害结果之间因果关系明确，如"未评估→数据泄露风险→用户损害"。'),

        ('⚠️ 潜在逻辑风险',
         '小型处理者标准可能被穿透监管，需持续关注关联交易认定标准变化。'),

        ('📋 持续监测要求',
         '1. 每月检索最高人民法院新案例\n'
         '2. 每季度更新法律条文变化\n'
         '3. 每年进行合规审计\n'
         '4. 建立政策预警机制')
    ]

    for check, content in check_items:
        p = doc.add_paragraph()
        p.add_run(check).bold = True
        doc.add_paragraph(content)

    # ===== 第七部分：律师建议 =====
    heading7 = doc.add_heading('第七部分：首席办案合伙人建议', level=1)

    p = doc.add_paragraph()
    p.add_run('基于上述穿透分析，提出以下战略建议：').bold = True

    suggestions = [
        '建议1：跨国公司应优先处理数据出境合规，预计2026年底前完成安全评估或认证。',
        '建议2：AI企业需在2026年7月15日前完成拟人化服务安全评估与算法备案。',
        '建议3：中小微企业应主动申请"小型个人信息处理者"资质，降低合规成本。',
        '建议4：平台企业需重构价格算法，避免违反《互联网平台价格行为规则》。',
        '建议5：建立"合规红利"思维，将网络安全标识认证转化为市场竞争力。',
        '建议6：开发"监管政策智能监测系统"，实现风险预警自动化。',
        '建议7：与国有数据服务商建立合作，分担跨境传输风险。',
        '建议8：设立AI伦理委员会，提前布局ESG（环境、社会、治理）评级。'
    ]

    for suggestion in suggestions:
        doc.add_paragraph(suggestion, style='List Bullet')

    # 保存文档
    desktop_path = os.path.expanduser("~/Desktop/中央网信办政策法律风险分析报告_2026.docx")
    doc.save(desktop_path)

    print(f"法律风险分析报告已生成：{desktop_path}")
    return desktop_path

if __name__ == "__main__":
    create_legal_risk_report()
