#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LegalRelationshipAnalyst 法律关系分析技能实现
分析法律文档，提取法律关系、证据链、风险点，生成专业Word报告
支持 --file, --text 参数
"""

import sys
import argparse
import os
import re
import json
import tempfile
import subprocess
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
import datetime
from dataclasses import dataclass

# 导入 PDF 解析库 (pymupdf)
try:
    import fitz  # pymupdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False
    print("警告: 未安装 pymupdf，将跳过PDF解析", file=sys.stderr)

# 导入 Word 文档解析库
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("警告: 未安装 python-docx，将跳过Word文档解析", file=sys.stderr)

# 导入 NLP 相关库
try:
    import jieba
    import jieba.posseg as pseg
    HAS_JIEBA = True
except ImportError:
    HAS_JIEBA = False
    print("警告: 未安装 jieba，将使用简单正则提取", file=sys.stderr)

# 导入报告生成库
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX_GEN = True
except ImportError:
    HAS_DOCX_GEN = False
    print("警告: 未安装 python-docx，无法生成Word报告", file=sys.stderr)


@dataclass
class LegalEntity:
    """法律实体"""
    name: str
    type: str  # "person", "company", "organization", "government"
    role: str  # "plaintiff", "defendant", "witness", "third_party"
    attributes: Dict[str, Any]

@dataclass
class LegalFact:
    """法律事实"""
    description: str
    timestamp: Optional[str]
    location: Optional[str]
    evidence_refs: List[str]
    legal_issues: List[str]

@dataclass
class Evidence:
    """证据"""
    id: str
    description: str
    type: str  # "document", "witness", "physical", "digital"
    source: str
    reliability: float  # 0.0-1.0
    issues: List[str]

@dataclass
class LegalRelationship:
    """法律关系"""
    from_entity: str
    to_entity: str
    relationship_type: str  # "contract", "debt", "ownership", "liability"
    content: str
    rights: List[str]
    obligations: List[str]

class LegalDocumentAnalyzer:
    """法律文档分析器"""

    def __init__(self):
        self.text = ""
        self.entities = []
        self.facts = []
        self.evidences = []
        self.relationships = []

        # 初始化jieba（如果可用）
        if HAS_JIEBA:
            # 添加法律专业词汇
            legal_terms = ["原告", "被告", "上诉人", "被上诉人", "法定代表人", "被执行人",
                          "债务人", "债权人", "担保人", "抵押人", "出质人", "质权人",
                          "用人单位", "劳动者", "出租人", "承租人", "出卖人", "买受人",
                          "许可人", "被许可人", "发行人", "认购人", "投保人", "保险人"]
            for term in legal_terms:
                jieba.add_word(term, freq=1000, tag='n')

    def load_file(self, file_path: str) -> bool:
        """加载文件内容"""
        if not os.path.exists(file_path):
            print(f"错误: 文件不存在: {file_path}", file=sys.stderr)
            return False

        ext = os.path.splitext(file_path)[1].lower()

        try:
            if ext == '.pdf' and HAS_PYPDF:
                self._load_pdf(file_path)
            elif ext in ['.docx', '.doc'] and HAS_DOCX:
                self._load_docx(file_path)
            elif ext in ['.txt', '.md', '.json', '.xml']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    self.text = f.read()
            else:
                print(f"错误: 不支持的文件格式: {ext}", file=sys.stderr)
                return False
            return True
        except Exception as e:
            print(f"读取文件错误: {e}", file=sys.stderr)
            return False

    def _load_pdf(self, file_path: str):
        """加载PDF文件"""
        text_parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                text = page.get_text()
                if text:
                    text_parts.append(text)
        self.text = '\n'.join(text_parts)

    def _load_docx(self, file_path: str):
        """加载Word文档"""
        doc = DocxDocument(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        self.text = '\n'.join(text_parts)

    def analyze(self):
        """分析文档内容"""
        if not self.text:
            print("错误: 没有文本内容可分析", file=sys.stderr)
            return

        print("开始分析法律文档...")

        # 提取实体
        self._extract_entities()

        # 提取事实
        self._extract_facts()

        # 提取证据
        self._extract_evidences()

        # 提取法律关系
        self._extract_relationships()

        print(f"分析完成: 找到{len(self.entities)}个实体, {len(self.facts)}个事实, "
              f"{len(self.evidences)}个证据, {len(self.relationships)}个法律关系")

    def _extract_entities(self):
        """提取法律实体"""
        if HAS_JIEBA:
            words = pseg.cut(self.text)
            current_entity = None

            for word, flag in words:
                if flag in ['nr', 'nz', 'nt', 'n'] and len(word) >= 2:
                    # 检查是否是法律实体
                    if self._is_legal_entity(word, flag):
                        entity_type = self._classify_entity_type(word, flag)
                        role = self._determine_entity_role(word)

                        self.entities.append(LegalEntity(
                            name=word,
                            type=entity_type,
                            role=role,
                            attributes={}
                        ))
        else:
            # 简单正则匹配
            patterns = [
                r'原告[：:]\s*([^\s，,。；;]+)',
                r'被告[：:]\s*([^\s，,。；;]+)',
                r'([\u4e00-\u9fff]{2,5}公司)',
                r'([\u4e00-\u9fff]{2,4}（[\u4e00-\u9fff]{2,4}）)',
            ]

            for pattern in patterns:
                matches = re.finditer(pattern, self.text)
                for match in matches:
                    name = match.group(1)
                    if name not in [e.name for e in self.entities]:
                        entity_type = self._classify_entity_type(name, 'n')
                        role = self._determine_entity_role(name)

                        self.entities.append(LegalEntity(
                            name=name,
                            type=entity_type,
                            role=role,
                            attributes={}
                        ))

    def _is_legal_entity(self, word: str, flag: str) -> bool:
        """判断是否为法律实体"""
        legal_keywords = ['公司', '法院', '检察院', '公安局', '事务所', '协会', '基金会']
        if any(keyword in word for keyword in legal_keywords):
            return True
        if flag == 'nr' and len(word) >= 2:  # 人名
            return True
        return False

    def _classify_entity_type(self, word: str, flag: str) -> str:
        """分类实体类型"""
        if '公司' in word or '企业' in word or '厂' in word:
            return 'company'
        elif '法院' in word or '检察院' in word or '公安局' in word:
            return 'government'
        elif '协会' in word or '基金会' in word or '组织' in word:
            return 'organization'
        elif flag == 'nr':
            return 'person'
        else:
            return 'organization'

    def _determine_entity_role(self, word: str) -> str:
        """确定实体角色"""
        if '原告' in word or '上诉人' in word:
            return 'plaintiff'
        elif '被告' in word or '被上诉人' in word:
            return 'defendant'
        elif '证人' in word:
            return 'witness'
        else:
            return 'third_party'

    def _extract_facts(self):
        """提取法律事实"""
        # 提取时间信息
        time_patterns = [
            r'(\d{4})年(\d{1,2})月(\d{1,2})日',
            r'(\d{4})-(\d{1,2})-(\d{1,2})',
            r'(\d{4})/(\d{1,2})/(\d{1,2})',
            r'于\s*(\d{4})年(\d{1,2})月(\d{1,2})日',
        ]

        sentences = re.split(r'[。！？；;]', self.text)
        for sentence in sentences:
            if len(sentence.strip()) < 10:
                continue

            # 查找时间
            timestamp = None
            for pattern in time_patterns:
                match = re.search(pattern, sentence)
                if match:
                    if pattern.startswith(r'(\d{4})年'):
                        timestamp = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2)}"
                    else:
                        timestamp = f"{match.group(1)}-{match.group(2).zfill(2)}-{match.group(3).zfill(2)}"
                    break

            # 查找金额
            amount_pattern = r'(?:人民币|金额|共计|总计)\s*[：:]?\s*([¥￥]?\d+(?:[,，]\d{3})*(?:\.\d+)?)\s*(?:元|万元|亿元)?'
            amount_match = re.search(amount_pattern, sentence)

            # 查找地点
            location_pattern = r'在\s*([\u4e00-\u9fff]{2,10}市|[\u4e00-\u9fff]{2,10}区|[\u4e00-\u9fff]{2,10}县)'
            location_match = re.search(location_pattern, sentence)

            if timestamp or amount_match or '约定' in sentence or '违约' in sentence:
                fact = LegalFact(
                    description=sentence.strip(),
                    timestamp=timestamp,
                    location=location_match.group(1) if location_match else None,
                    evidence_refs=[],
                    legal_issues=[]
                )
                self.facts.append(fact)

    def _extract_evidences(self):
        """提取证据"""
        evidence_keywords = ['合同', '协议', '借条', '收据', '发票', '银行流水', '聊天记录',
                           '录音', '录像', '照片', '鉴定意见', '证人证言', '书证', '物证']

        sentences = re.split(r'[。！？；;]', self.text)
        evidence_id = 1

        for sentence in sentences:
            for keyword in evidence_keywords:
                if keyword in sentence:
                    # 提取证据描述
                    desc = sentence.strip()

                    # 判断证据类型
                    evidence_type = "document"
                    if keyword in ['录音', '录像', '照片']:
                        evidence_type = "digital"
                    elif keyword in ['证人证言']:
                        evidence_type = "witness"
                    elif keyword in ['物证']:
                        evidence_type = "physical"

                    # 评估可靠性
                    reliability = 0.7
                    if '原件' in sentence:
                        reliability = 0.9
                    elif '复印件' in sentence:
                        reliability = 0.5

                    evidence = Evidence(
                        id=f"E{evidence_id}",
                        description=desc,
                        type=evidence_type,
                        source="文档中提及",
                        reliability=reliability,
                        issues=[]
                    )
                    self.evidences.append(evidence)
                    evidence_id += 1
                    break

    def _extract_relationships(self):
        """提取法律关系"""
        if len(self.entities) < 2:
            return

        # 查找实体之间的关系
        relationship_keywords = ['与', '和', '向', '对', '欠', '借', '租', '卖', '买', '许可', '授权']

        for i, entity1 in enumerate(self.entities):
            for j, entity2 in enumerate(self.entities):
                if i >= j:
                    continue

                # 查找两个实体同时出现的句子
                pattern = f"{re.escape(entity1.name)}.*?{re.escape(entity2.name)}|{re.escape(entity2.name)}.*?{re.escape(entity1.name)}"
                matches = re.finditer(pattern, self.text, re.DOTALL)

                for match in matches:
                    context = match.group(0)
                    if any(keyword in context for keyword in relationship_keywords):
                        # 确定关系类型
                        rel_type = "unknown"
                        if '欠' in context or '借' in context:
                            rel_type = "debt"
                        elif '租' in context:
                            rel_type = "lease"
                        elif '卖' in context or '买' in context:
                            rel_type = "sale"
                        elif '许可' in context or '授权' in context:
                            rel_type = "license"
                        elif '雇佣' in context or '聘用' in context:
                            rel_type = "employment"

                        relationship = LegalRelationship(
                            from_entity=entity1.name,
                            to_entity=entity2.name,
                            relationship_type=rel_type,
                            content=context[:100] + "..." if len(context) > 100 else context,
                            rights=[],
                            obligations=[]
                        )
                        self.relationships.append(relationship)

    def generate_report(self, output_path: str = None) -> str:
        """生成Word报告"""
        if not HAS_DOCX_GEN:
            print("错误: 需要python-docx库来生成报告", file=sys.stderr)
            return ""

        if output_path is None:
            desktop = Path.home() / "Desktop"
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = str(desktop / f"法律关系分析报告_{timestamp}.docx")

        doc = Document()

        # 标题
        title = doc.add_heading('法律关系分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 报告信息
        doc.add_paragraph(f"生成时间: {datetime.datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}")
        doc.add_paragraph(f"分析文档: {getattr(self, 'source_file', '文本输入')}")
        doc.add_paragraph()

        # 1. 法律实体分析
        doc.add_heading('一、法律主体分析', level=1)
        if self.entities:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '主体名称'
            hdr_cells[1].text = '类型'
            hdr_cells[2].text = '角色'
            hdr_cells[3].text = '备注'

            for entity in self.entities:
                row_cells = table.add_row().cells
                row_cells[0].text = entity.name
                row_cells[1].text = entity.type
                row_cells[2].text = entity.role
                row_cells[3].text = str(entity.attributes)
        else:
            doc.add_paragraph('未识别到明确的法律主体。')

        doc.add_paragraph()

        # 2. 法律事实脉络
        doc.add_heading('二、事实脉络表', level=1)
        if self.facts:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '时间'
            hdr_cells[1].text = '事实描述'
            hdr_cells[2].text = '地点'
            hdr_cells[3].text = '关联证据'

            for fact in self.facts:
                row_cells = table.add_row().cells
                row_cells[0].text = fact.timestamp or '未知'
                row_cells[1].text = fact.description[:100] + "..." if len(fact.description) > 100 else fact.description
                row_cells[2].text = fact.location or '未知'
                row_cells[3].text = ', '.join(fact.evidence_refs) if fact.evidence_refs else '无'
        else:
            doc.add_paragraph('未提取到明确的法律事实。')

        doc.add_paragraph()

        # 3. 证据链分析
        doc.add_heading('三、证据链分析', level=1)
        if self.evidences:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '证据编号'
            hdr_cells[1].text = '证据描述'
            hdr_cells[2].text = '类型'
            hdr_cells[3].text = '可靠性'
            hdr_cells[4].text = '问题点'

            for evidence in self.evidences:
                row_cells = table.add_row().cells
                row_cells[0].text = evidence.id
                row_cells[1].text = evidence.description[:80] + "..." if len(evidence.description) > 80 else evidence.description
                row_cells[2].text = evidence.type
                row_cells[3].text = f"{evidence.reliability*100:.0f}%"
                row_cells[4].text = ', '.join(evidence.issues) if evidence.issues else '无'
        else:
            doc.add_paragraph('未识别到明确的证据。')

        doc.add_paragraph()

        # 4. 法律关系图谱
        doc.add_heading('四、法律关系图谱', level=1)
        if self.relationships:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '主体A'
            hdr_cells[1].text = '主体B'
            hdr_cells[2].text = '关系类型'
            hdr_cells[3].text = '关系内容'
            hdr_cells[4].text = '权利义务'

            for rel in self.relationships:
                row_cells = table.add_row().cells
                row_cells[0].text = rel.from_entity
                row_cells[1].text = rel.to_entity
                row_cells[2].text = rel.relationship_type
                row_cells[3].text = rel.content[:80] + "..." if len(rel.content) > 80 else rel.content
                row_cells[4].text = f"权利: {len(rel.rights)}项, 义务: {len(rel.obligations)}项"
        else:
            doc.add_paragraph('未识别到明确的法律关系。')

        doc.add_paragraph()

        # 5. 风险红绿灯
        doc.add_heading('五、风险红绿灯', level=1)

        # 风险分析
        risks = []

        # 证据风险
        low_reliability_evidences = [e for e in self.evidences if e.reliability < 0.6]
        if low_reliability_evidences:
            risks.append({
                'level': '高风险',
                'description': f"发现{len(low_reliability_evidences)}个低可靠性证据",
                'suggestion': '建议补充原件或加强证据链'
            })

        # 主体风险
        if len(self.entities) < 2:
            risks.append({
                'level': '中风险',
                'description': '法律主体不足，可能影响法律关系认定',
                'suggestion': '建议明确所有相关方的主体资格'
            })

        # 事实风险
        incomplete_facts = [f for f in self.facts if not f.timestamp]
        if incomplete_facts:
            risks.append({
                'level': '中风险',
                'description': f"发现{len(incomplete_facts)}个事实缺少明确时间点",
                'suggestion': '建议补充时间证据'
            })

        if not risks:
            risks.append({
                'level': '低风险',
                'description': '未发现明显高风险项',
                'suggestion': '建议继续完善证据材料'
            })

        # 风险表格
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '风险等级'
        hdr_cells[1].text = '风险描述'
        hdr_cells[2].text = '应对建议'

        for risk in risks:
            row_cells = table.add_row().cells
            row_cells[0].text = risk['level']
            # 根据风险等级设置颜色
            if risk['level'] == '高风险':
                run = row_cells[0].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 0, 0)
            elif risk['level'] == '中风险':
                run = row_cells[0].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(255, 165, 0)
            else:
                run = row_cells[0].paragraphs[0].runs[0]
                run.font.color.rgb = RGBColor(0, 128, 0)

            row_cells[1].text = risk['description']
            row_cells[2].text = risk['suggestion']

        doc.add_paragraph()

        # 6. 应对建议
        doc.add_heading('六、应对建议', level=1)
        doc.add_paragraph('1. 证据补强建议：')
        doc.add_paragraph('   - 收集并固定所有证据原件')
        doc.add_paragraph('   - 对关键证据进行公证或鉴定')
        doc.add_paragraph('   - 补充时间戳、地点等关键信息')

        doc.add_paragraph('2. 法律策略建议：')
        doc.add_paragraph('   - 明确诉讼主体资格')
        doc.add_paragraph('   - 确定管辖法院')
        doc.add_paragraph('   - 评估诉讼时效风险')

        doc.add_paragraph('3. 谈判和解建议：')
        doc.add_paragraph('   - 基于现有证据评估和解可能性')
        doc.add_paragraph('   - 制定阶梯式谈判方案')
        doc.add_paragraph('   - 准备调解或仲裁备选方案')

        # 保存文档
        doc.save(output_path)
        print(f"报告已生成: {output_path}")
        return output_path


def main():
    parser = argparse.ArgumentParser(description='法律关系分析专家系统')
    parser.add_argument('--file', type=str, help='要分析的文件路径')
    parser.add_argument('--text', type=str, help='直接分析文本内容')
    parser.add_argument('--output', type=str, help='输出报告路径（可选）')

    args = parser.parse_args()

    if not args.file and not args.text:
        parser.print_help()
        print("\n示例:")
        print("  python legal_analyst.py --file 合同.pdf")
        print("  python legal_analyst.py --text '原告张三诉被告李四借款纠纷...'")
        return 1

    analyzer = LegalDocumentAnalyzer()

    if args.file:
        if not analyzer.load_file(args.file):
            return 1
        analyzer.source_file = args.file
    elif args.text:
        analyzer.text = args.text
        analyzer.source_file = "文本输入"

    analyzer.analyze()

    output_path = analyzer.generate_report(args.output)

    if output_path and os.path.exists(output_path):
        print(f"\n分析完成！报告已保存至: {output_path}")
        print("\n报告包含以下章节：")
        print("1. 法律主体分析")
        print("2. 事实脉络表")
        print("3. 证据链分析")
        print("4. 法律关系图谱")
        print("5. 风险红绿灯")
        print("6. 应对建议")

    return 0


if __name__ == '__main__':
    sys.exit(main())
