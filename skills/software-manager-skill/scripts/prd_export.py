#!/usr/bin/env python3
"""
Export PRD markdown to DOCX with Mermaid diagrams rendered as PNG images.
Usage: python3 prd_export.py <input.md> <output.docx>
Requires: python-docx, PIL (for image size), node+ puppeteer (for mermaid)
"""
import sys
import re
import os
import subprocess
import tempfile
import shutil
from pathlib import Path

def _find_mermaid_script() -> str:
    """Find mermaid_render_multi.js, checking multiple possible locations."""
    import os as _os
    # 尝试的路径：优先 __file__ 所在目录，再从 cwd 和常见 Codex skill 目录推导。
    _search_dirs = []
    if '__file__' in dir():
        _search_dirs.append(str(Path(__file__).resolve().parent))
    # 从 cwd 推导（兼容 cd && python 方式调用）
    _search_dirs.append(str(Path.cwd()))
    _search_dirs.append(str(Path.cwd().parent))  # 父目录（scripts/ → skill root）
    # 常见 skill 安装路径
    _search_dirs.extend([
        str(Path.home() / '.codex/skills/software-manager-skill/scripts'),
        str(Path.home() / '.agents/skills/software-manager-skill/scripts'),
    ])
    for _dir in _search_dirs:
        _candidate = _os.path.join(_dir, 'mermaid_render_multi.js')
        if _os.path.isfile(_candidate):
            return _candidate
    return ''

def render_mermaid_to_png(mermaid_code: str, output_path: str) -> bool:
    """Render a mermaid diagram to PNG using node+puppeteer (multi-browser).
    Supports Chrome, Edge, Firefox on Windows/Mac/Linux. Auto-detects available browser.
    """
    # Write mermaid input to temp file
    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, dir='/tmp') as f:
        f.write(mermaid_code)
        mmd_path = f.name

    try:
        # 查找 mermaid_render_multi.js
        mm_script = _find_mermaid_script()
        if not mm_script or not os.path.isfile(mm_script):
            print("ERROR: mermaid_render_multi.js not found in any search path.", file=sys.stderr)
            os.unlink(mmd_path)
            return False

        result = subprocess.run(
            ['node', mm_script, mmd_path, output_path],
            capture_output=True, timeout=60
        )
        return result.returncode == 0 and os.path.exists(output_path)
    finally:
        try:
            os.unlink(mmd_path)
        except:
            pass


def parse_markdown_sections(text: str) -> list:
    """Parse markdown into sections with content types."""
    sections = []
    # Split on mermaid blocks first
    parts = re.split(r'(```mermaid\n[\s\S]*?```)', text)

    for part in parts:
        if part.startswith('```mermaid'):
            # Extract mermaid code (remove ```mermaid and ```)
            code = re.sub(r'^```mermaid\n?', '', part)
            code = re.sub(r'\n?```$', '', code)
            sections.append({'type': 'mermaid', 'content': code.strip()})
        elif part.strip():
            sections.append({'type': 'markdown', 'content': part})

    return sections


def heading_level(line: str) -> int:
    """Return heading level (1-6) or 0 if not a heading."""
    m = re.match(r'^(#{1,6})\s+', line)
    return len(m.group(1)) if m else 0


def export_md_to_docx(md_text: str, docx_path: str, title: str = "PRD Document"):
    """Export markdown text with mermaid to DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip3 install python-docx", file=sys.stderr)
        return False

    # Create a temp dir for mermaid images
    img_dir = tempfile.mkdtemp(prefix='prd_mmd_')
    img_counter = [0]

    def next_img_path() -> str:
        img_counter[0] += 1
        return os.path.join(img_dir, f'mermaid_{img_counter[0]:03d}.png')

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add the full content (markdown + mermaid sections processed)
    lines = md_text.split('\n')
    in_code_block = False
    code_block_content = []
    code_lang = ''

    def flush_code_block(doc, lang, content_lines):
        if not content_lines:
            return
        code_text = '\n'.join(content_lines)
        if lang == 'mermaid':
            # Render mermaid to image
            img_path = next_img_path()
            ok = render_mermaid_to_png(code_text, img_path)
            if ok and os.path.exists(img_path):
                # Add image to doc
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(img_path, width=Inches(5.5))
                doc.add_paragraph()  # spacing
            else:
                # Fallback: add mermaid code as monospace text with error note
                para = doc.add_paragraph()
                para.style = 'No Spacing'
                run = para.add_run('[Mermaid图表渲染失败，已降级为代码块]\n' + code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(9)
        else:
            # Add as monospace text
            para = doc.add_paragraph()
            para.style = 'No Spacing'
            run = para.add_run(code_text)
            run.font.name = 'Courier New'
            run.font.size = Pt(9)

    for line in lines:
        if line.startswith('```'):
            if not in_code_block:
                # Start of code block
                in_code_block = True
                code_lang = line[3:].strip()
                code_block_content = []
            else:
                # End of code block
                in_code_block = False
                flush_code_block(doc, code_lang, code_block_content)
                code_lang = ''
                code_block_content = []
        elif in_code_block:
            code_block_content.append(line)
        else:
            # Regular markdown line
            hlvl = heading_level(line)
            if hlvl > 0:
                # Remove the # markers for heading
                heading_text = line.lstrip('#').strip()
                doc.add_heading(heading_text, level=hlvl)
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                doc.add_paragraph(line[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s+', line):
                # Numbered list
                m = re.match(r'^(\d+\.)\s+(.*)', line)
                if m:
                    p = doc.add_paragraph(m.group(2), style='List Number')
                    p.style._element.set(qn('w:numId'), m.group(1))
            elif line.startswith('|'):
                # Table-like line - skip individual lines, add as plain text
                # (python-docx table parsing is complex, simplified here)
                doc.add_paragraph(line)
            elif line.strip() == '':
                doc.add_paragraph()
            else:
                # Plain paragraph - handle inline code
                inline_code = re.sub(r'`([^`]+)`', r'\1', line)
                doc.add_paragraph(inline_code)

    # Save
    doc.save(docx_path)

    # Cleanup temp images
    try:
        shutil.rmtree(img_dir)
    except:
        pass

    return True


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python3 prd_export.py <input.md> <output.docx> [title]")
        sys.exit(1)

    # ── 路径预处理：清理 Markdown 链接语法，并统一 Windows → WSL 路径 ──
    def normalize_path(p: str) -> str:
        """清理文件名中的 Markdown 链接语法，并转换 Windows 路径为 WSL 路径。"""
        import re
        # 清理 Markdown 链接语法：`[文件名](url)` → `文件名`
        cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', p)
        # 清理尾部 URL（裸 URL 被当成文件名的一部分）
        cleaned = re.sub(r'https?://\S+$', '', cleaned).strip()
        cleaned = re.sub(r'http://\S+$', '', cleaned).strip()
        # Windows 路径转换（WSL 环境）
        if cleaned.startswith('/mnt/c/') or cleaned.startswith('C:\\') or cleaned.startswith('C:/'):
            # already WSL or Windows format - let Path handle it
            pass
        elif len(cleaned) > 2 and cleaned[1] == ':':
            # Windows absolute path like C:\Users\...
            drive = cleaned[0].upper()
            rest = cleaned[2:].replace('\\', '/').replace(':', '')
            cleaned = f'/mnt/{drive.lower()}/{rest}'
        return cleaned

    md_path_raw = sys.argv[1]
    docx_path_raw = sys.argv[2]

    md_path = normalize_path(md_path_raw)
    docx_path = normalize_path(docx_path_raw)
    title = sys.argv[3] if len(sys.argv) > 3 else Path(md_path).stem

    if not os.path.exists(md_path):
        print(f"ERROR: Input file not found: {md_path}  (raw: {md_path_raw})", file=sys.stderr)
        sys.exit(1)

    md_text = Path(md_path).read_text(encoding='utf-8')
    success = export_md_to_docx(md_text, docx_path, title)

    if success:
        print(f"OK: {docx_path}")
    else:
        print(f"ERROR: Export failed", file=sys.stderr)
        sys.exit(1)
