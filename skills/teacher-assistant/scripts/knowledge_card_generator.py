#!/usr/bin/env python3
"""
教学知识卡片生成器 v2.0
生成精美的可打印知识卡片，支持 HTML + Word(.docx) 双格式输出。
HTML 版本提供最佳视觉效果，可直接用浏览器打印/另存为PDF。

用法：
  python scripts/knowledge_card_generator.py <input.md> [output_prefix]
  python scripts/knowledge_card_generator.py --text "<markdown内容>" [output_prefix]
"""

import re
import sys
import subprocess
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── 颜色配置（现代设计风格） ─────────────────────────────────────────────
PALETTE = {
    "primary":      "#2563EB",   # 主色调蓝
    "primary_dark": "#1D4ED8",   # 深蓝
    "primary_light":"#EFF6FF",   # 极浅蓝
    "accent":       "#7C3AED",   # 紫色强调
    "success":      "#059669",   # 绿色
    "warning":      "#D97706",   # 橙黄
    "danger":       "#DC2626",   # 红色
    "text_dark":    "#111827",   # 深色正文
    "text_mid":     "#374151",   # 中灰正文
    "text_light":   "#6B7280",   # 浅灰辅助
    "border":       "#E5E7EB",   # 边框灰
    "bg_page":      "#F3F4F6",   # 页面背景
}

SECTION_STYLES = {
    "核心定义":     {"icon": "📌", "bg": "#EFF6FF", "border": "#2563EB", "title_color": "#1E40AF"},
    "关键公式":     {"icon": "📐", "bg": "#F0FDF4", "border": "#059669", "title_color": "#065F46"},
    "关键规律":     {"icon": "📐", "bg": "#F0FDF4", "border": "#059669", "title_color": "#065F46"},
    "典型例题":     {"icon": "🌰", "bg": "#FFFBEB", "border": "#D97706", "title_color": "#92400E"},
    "记忆口诀":     {"icon": "💡", "bg": "#FDF4FF", "border": "#7C3AED", "title_color": "#5B21B6"},
    "易错提示":     {"icon": "⚠️", "bg": "#FEF2F2", "border": "#DC2626", "title_color": "#991B1B"},
    "相关知识":     {"icon": "🔗", "bg": "#F8FAFC", "border": "#64748B", "title_color": "#334155"},
    "default":      {"icon": "📝", "bg": "#FAFAFA",  "border": "#9CA3AF", "title_color": "#374151"},
}

FONT_SONG  = "宋体"
FONT_HEI   = "黑体"
FONT_KAITI = "楷体"


# ══════════════════════════════════════════════════════════════════
#   Markdown 解析
# ══════════════════════════════════════════════════════════════════

def parse_cards(md_text: str) -> list[dict]:
    """将 Markdown 文本按 '## ...' 切分为卡片列表"""
    blocks = re.split(r'\n(?=## )', md_text.strip())
    cards = []
    for block in blocks:
        lines = block.strip().split('\n')
        if not lines:
            continue
        header_line = lines[0]
        title_match = re.search(r'知识卡片[：:]\s*(.+)', header_line)
        if not title_match:
            title_match = re.match(r'#{1,3}\s*(.+)', header_line)
        title = title_match.group(1).strip() if title_match else header_line.strip('# ').strip()
        # 去掉 emoji 前缀（如 📘）
        title = re.sub(r'^[^\w\s（(【\[]+\s*', '', title).strip()

        meta_line = ''
        body_start = 1
        for i, line in enumerate(lines[1:], 1):
            if line.strip().startswith('>'):
                meta_line = line.strip().lstrip('> ').strip()
                body_start = i + 1
                break

        subject, grade, index_str = '', '', ''
        if meta_line:
            for seg in re.split(r'[|｜]', meta_line):
                seg = seg.strip()
                if '学科' in seg:
                    subject = re.sub(r'学科[：:]\s*', '', seg).strip()
                elif '年级' in seg:
                    grade = re.sub(r'年级[：:]\s*', '', seg).strip()
                elif '编号' in seg:
                    index_str = re.sub(r'编号[：:]\s*', '', seg).strip()

        body_lines = lines[body_start:]
        sections = []
        current_heading, current_heading_raw, current_content = '', '', []
        for line in body_lines:
            h_match = re.match(r'^#{2,4}\s+(.+)', line)
            if h_match:
                if current_heading or current_content:
                    sections.append((current_heading, current_heading_raw,
                                     '\n'.join(current_content).strip()))
                raw_h = h_match.group(1).strip()
                clean_h = re.sub(r'^[^\w\s（(【\[]+\s*', '', raw_h).strip()
                current_heading = clean_h
                current_heading_raw = raw_h
                current_content = []
            elif re.match(r'^---+$', line.strip()):
                continue
            else:
                current_content.append(line)
        if current_heading or current_content:
            sections.append((current_heading, current_heading_raw,
                             '\n'.join(current_content).strip()))

        cards.append({
            'title':    title,
            'subject':  subject,
            'grade':    grade,
            'index':    index_str,
            'sections': sections,
        })
    return cards


def _get_section_style(heading: str) -> dict:
    for key, style in SECTION_STYLES.items():
        if key != "default" and key in heading:
            return style
    return SECTION_STYLES["default"]


# ══════════════════════════════════════════════════════════════════
#   HTML 生成（主要输出格式，视觉最佳）
# ══════════════════════════════════════════════════════════════════

def _md_inline_to_html(text: str) -> str:
    """简单内联 Markdown → HTML（粗体、代码）"""
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)
    text = re.sub(r'\*(.+?)\*',     r'<em>\1</em>',         text)
    text = re.sub(r'`(.+?)`',       r'<code>\1</code>',      text)
    return text


def _render_content_html(content: str, section_style: dict) -> str:
    """将节内容渲染为 HTML"""
    html_parts = []
    for line in content.split('\n'):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith('>'):
            line_text = _md_inline_to_html(line.lstrip('> ').strip())
            html_parts.append(f'<blockquote class="quote">{line_text}</blockquote>')
        elif re.match(r'^\s*[-*•]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            text = re.sub(r'^\s*[-*•\d.]+\s+', '', line).strip()
            text = _md_inline_to_html(text)
            is_numbered = bool(re.match(r'^\s*\d+\.', line))
            num = re.match(r'^\s*(\d+)\.', line)
            if is_numbered and num:
                html_parts.append(f'<div class="list-item numbered"><span class="num">{num.group(1)}</span><span>{text}</span></div>')
            else:
                html_parts.append(f'<div class="list-item"><span class="bullet">•</span><span>{text}</span></div>')
        else:
            line_html = _md_inline_to_html(line)
            html_parts.append(f'<p class="content-line">{line_html}</p>')
    return '\n'.join(html_parts)


def _render_card_html(card: dict) -> str:
    """将一张卡片渲染为 HTML 片段"""
    meta_parts = []
    if card['subject']: meta_parts.append(f'<span class="meta-tag">📚 {card["subject"]}</span>')
    if card['grade']:   meta_parts.append(f'<span class="meta-tag">🎓 {card["grade"]}</span>')
    if card['index']:   meta_parts.append(f'<span class="meta-tag">🔢 No.{card["index"]}</span>')
    meta_html = f'<div class="card-meta">{"".join(meta_parts)}</div>' if meta_parts else ''

    sections_html = []
    for heading, heading_raw, content in card['sections']:
        if not heading and not content:
            continue
        style = _get_section_style(heading)
        icon = style['icon']
        bg = style['bg']
        border = style['border']
        title_color = style['title_color']
        content_html = _render_content_html(content, style)
        sections_html.append(f'''
        <div class="section" style="background:{bg};border-left:4px solid {border}">
            <div class="section-title" style="color:{title_color}">
                <span class="section-icon">{icon}</span>
                {heading}
            </div>
            <div class="section-body">{content_html}</div>
        </div>''')

    return f'''
<div class="knowledge-card">
    <div class="card-header">
        <div class="card-title">📘 {card['title']}</div>
        {meta_html}
    </div>
    <div class="card-body">
        {''.join(sections_html)}
    </div>
    <div class="card-footer">
        <span>✦ 教学知识卡片</span>
        <span>{datetime.now().strftime("%Y.%m.%d")}</span>
    </div>
</div>'''


def generate_html(cards: list[dict], output_path: str) -> Path:
    """生成高颜值 HTML 知识卡片文件"""
    cards_html = '\n'.join(_render_card_html(c) for c in cards)
    cards_count = len(cards)

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>教学知识卡片</title>
<style>
  /* ── 基础重置 ── */
  *, *::before, *::after {{ box-sizing: border-box; margin: 0; padding: 0; }}
  html {{ font-size: 14px; }}
  body {{
    font-family: "Microsoft YaHei", "PingFang SC", "Helvetica Neue", sans-serif;
    background: #F0F4F8;
    color: #1F2937;
    padding: 32px 20px;
    min-height: 100vh;
  }}

  /* ── 页头 ── */
  .page-header {{
    text-align: center;
    margin-bottom: 36px;
  }}
  .page-header h1 {{
    font-size: 1.8rem;
    font-weight: 700;
    color: #1E3A8A;
    letter-spacing: 0.05em;
    display: inline-flex;
    align-items: center;
    gap: 10px;
  }}
  .page-header .subtitle {{
    font-size: 0.85rem;
    color: #6B7280;
    margin-top: 6px;
  }}
  .page-header .badge {{
    display: inline-block;
    background: #DBEAFE;
    color: #1D4ED8;
    font-size: 0.75rem;
    padding: 2px 10px;
    border-radius: 999px;
    margin-top: 8px;
  }}

  /* ── 卡片容器 ── */
  .cards-grid {{
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(520px, 1fr));
    gap: 28px;
    max-width: 1120px;
    margin: 0 auto;
  }}

  /* ── 单张知识卡片 ── */
  .knowledge-card {{
    background: #fff;
    border-radius: 16px;
    box-shadow: 0 4px 6px -1px rgba(0,0,0,.07), 0 10px 15px -3px rgba(0,0,0,.05);
    overflow: hidden;
    transition: box-shadow 0.2s;
    break-inside: avoid;
    page-break-inside: avoid;
  }}
  .knowledge-card:hover {{
    box-shadow: 0 8px 25px -5px rgba(37,99,235,.15), 0 4px 6px -2px rgba(0,0,0,.05);
  }}

  /* ── 卡片头部（渐变） ── */
  .card-header {{
    background: linear-gradient(135deg, #1D4ED8 0%, #7C3AED 100%);
    padding: 18px 20px 14px;
    position: relative;
    overflow: hidden;
  }}
  .card-header::before {{
    content: '';
    position: absolute;
    top: -20px; right: -20px;
    width: 100px; height: 100px;
    background: rgba(255,255,255,0.08);
    border-radius: 50%;
  }}
  .card-header::after {{
    content: '';
    position: absolute;
    bottom: -30px; left: 30%;
    width: 150px; height: 150px;
    background: rgba(255,255,255,0.04);
    border-radius: 50%;
  }}
  .card-title {{
    font-size: 1.2rem;
    font-weight: 700;
    color: #fff;
    letter-spacing: 0.02em;
    position: relative;
    z-index: 1;
  }}
  .card-meta {{
    margin-top: 8px;
    display: flex;
    gap: 8px;
    flex-wrap: wrap;
    position: relative;
    z-index: 1;
  }}
  .meta-tag {{
    background: rgba(255,255,255,0.2);
    color: rgba(255,255,255,0.95);
    font-size: 0.75rem;
    padding: 2px 10px;
    border-radius: 999px;
    backdrop-filter: blur(4px);
  }}

  /* ── 卡片正文 ── */
  .card-body {{
    padding: 16px 16px 4px;
    display: flex;
    flex-direction: column;
    gap: 10px;
  }}

  /* ── 各节 ── */
  .section {{
    border-radius: 10px;
    padding: 12px 14px;
    border-left-width: 4px;
    border-left-style: solid;
  }}
  .section-title {{
    font-size: 0.88rem;
    font-weight: 700;
    margin-bottom: 8px;
    display: flex;
    align-items: center;
    gap: 6px;
    letter-spacing: 0.03em;
  }}
  .section-icon {{ font-size: 1rem; }}
  .section-body {{ font-size: 0.85rem; line-height: 1.7; color: #374151; }}

  /* ── 正文元素 ── */
  .content-line {{ margin-bottom: 4px; }}
  .list-item {{
    display: flex;
    align-items: flex-start;
    gap: 6px;
    margin-bottom: 4px;
  }}
  .bullet {{
    color: #60A5FA;
    font-weight: 700;
    flex-shrink: 0;
    margin-top: 1px;
  }}
  .num {{
    background: #2563EB;
    color: #fff;
    font-size: 0.7rem;
    font-weight: 700;
    min-width: 18px;
    height: 18px;
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
    margin-top: 2px;
  }}
  .quote {{
    border-left: 3px solid #A78BFA;
    padding: 6px 10px;
    background: rgba(167,139,250,0.08);
    border-radius: 0 6px 6px 0;
    color: #5B21B6;
    font-style: italic;
    margin: 4px 0;
    font-size: 0.83rem;
  }}
  code {{
    background: #F1F5F9;
    border: 1px solid #E2E8F0;
    border-radius: 4px;
    padding: 1px 5px;
    font-family: "Courier New", monospace;
    font-size: 0.82em;
    color: #0F172A;
  }}
  strong {{ color: #1E40AF; font-weight: 700; }}

  /* ── 卡片底部 ── */
  .card-footer {{
    display: flex;
    justify-content: space-between;
    align-items: center;
    padding: 8px 16px 12px;
    font-size: 0.7rem;
    color: #9CA3AF;
    border-top: 1px solid #F3F4F6;
    margin-top: 8px;
  }}

  /* ── 打印样式 ── */
  @media print {{
    body {{
      background: #fff;
      padding: 0;
    }}
    .page-header {{ margin-bottom: 20px; }}
    .cards-grid {{
      grid-template-columns: 1fr 1fr;
      gap: 16px;
      max-width: 100%;
    }}
    .knowledge-card {{
      box-shadow: 0 0 0 1px #E5E7EB;
      border-radius: 10px;
    }}
    @page {{
      size: A4;
      margin: 15mm 12mm;
    }}
  }}

  /* ── 响应式 ── */
  @media (max-width: 600px) {{
    .cards-grid {{ grid-template-columns: 1fr; }}
  }}
</style>
</head>
<body>
<div class="page-header">
  <h1>📚 教学知识卡片</h1>
  <div class="subtitle">Teaching Knowledge Cards · {datetime.now().strftime("%Y 年 %m 月 %d 日")}</div>
  <div class="badge">共 {cards_count} 张卡片</div>
</div>
<div class="cards-grid">
{cards_html}
</div>
</body>
</html>'''

    out_path = Path(output_path)
    out_path.write_text(html, encoding='utf-8')
    print(f"[✓] HTML 知识卡片已保存：{out_path}")
    return out_path


# ══════════════════════════════════════════════════════════════════
#   Word 生成（美化版）
# ══════════════════════════════════════════════════════════════════

def _set_shading(cell, color_hex: str):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _run_font(run, name=FONT_SONG, size=10, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass
    if color:
        color = color.lstrip('#')
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _add_card_to_doc(doc, card: dict):
    """在 doc 中添加一张精美知识卡片"""
    # 外层 1×1 表格作卡片边框
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    try:
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    tbl.columns[0].width = Cm(15.5)
    cell = tbl.cell(0, 0)
    cell._tc.get_or_add_tcPr()

    # ── 大标题（渐变模拟：深蓝背景） ──────────────────────────
    _set_shading(cell, "1D4ED8")
    title_p = cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = title_p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after  = Pt(3)
    run = title_p.add_run(f"📘 {card['title']}")
    _run_font(run, name=FONT_HEI, size=14, bold=True, color="#FFFFFF")

    # 元信息行
    meta_parts = []
    if card['subject']: meta_parts.append(f"📚 {card['subject']}")
    if card['grade']:   meta_parts.append(f"🎓 {card['grade']}")
    if card['index']:   meta_parts.append(f"No.{card['index']}")
    if meta_parts:
        meta_p = cell.add_paragraph()
        meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_p.paragraph_format.space_before = Pt(0)
        meta_p.paragraph_format.space_after  = Pt(8)
        run_m = meta_p.add_run("  ·  ".join(meta_parts))
        _run_font(run_m, name=FONT_KAITI, size=9, color="#BFDBFE")
        _set_shading(cell, "1D4ED8")

    # ── 各节内容 ──────────────────────────────────────────────
    DOCX_SECTION_STYLES = {
        "核心定义":  ("EFF6FF", "1E40AF"),
        "关键公式":  ("F0FDF4", "065F46"),
        "关键规律":  ("F0FDF4", "065F46"),
        "典型例题":  ("FFFBEB", "92400E"),
        "记忆口诀":  ("FDF4FF", "5B21B6"),
        "易错提示":  ("FEF2F2", "991B1B"),
        "相关知识":  ("F8FAFC", "334155"),
    }
    DEFAULT_DOCX_STYLE = ("FAFAFA", "374151")

    for heading, heading_raw, content in card['sections']:
        if not heading and not content:
            continue

        bg, hcolor = DEFAULT_DOCX_STYLE
        for key, val in DOCX_SECTION_STYLES.items():
            if key in heading:
                bg, hcolor = val
                break

        # 节背景颜色用独立段落模拟（Word 中每段单独着色）
        h_p = cell.add_paragraph()
        h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_p.paragraph_format.space_before = Pt(8)
        h_p.paragraph_format.space_after  = Pt(3)
        h_p.paragraph_format.left_indent  = Cm(0.3)
        run_h = h_p.add_run(f"▌ {heading}")
        _run_font(run_h, name=FONT_HEI, size=10.5, bold=True, color=f"#{hcolor}")

        if content:
            for line in content.split('\n'):
                line = line.rstrip()
                if not line:
                    continue
                if line.startswith('>'):
                    line = line.lstrip('> ').strip()
                    c_p = cell.add_paragraph()
                    c_p.paragraph_format.left_indent  = Cm(0.8)
                    c_p.paragraph_format.space_before = Pt(1)
                    c_p.paragraph_format.space_after  = Pt(2)
                    run_c = c_p.add_run(f"❝  {line}  ❞")
                    _run_font(run_c, name=FONT_KAITI, size=9, italic=True, color="#5B21B6")
                elif re.match(r'^\s*[-*•]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
                    is_numbered = bool(re.match(r'^\s*\d+\.', line))
                    num = re.match(r'^\s*(\d+)\.', line)
                    text = re.sub(r'^\s*[-*•\d.]+\s+', '', line).strip()
                    c_p = cell.add_paragraph()
                    c_p.paragraph_format.left_indent  = Cm(0.8)
                    c_p.paragraph_format.space_before = Pt(1)
                    c_p.paragraph_format.space_after  = Pt(2)
                    prefix = f"  {num.group(1)}." if is_numbered and num else "  •"
                    parts = re.split(r'\*\*(.+?)\*\*', text)
                    first = True
                    for i, part in enumerate(parts):
                        if not part:
                            continue
                        run_c = c_p.add_run((f"{prefix} " if first else "") + part)
                        first = False
                        _run_font(run_c, name=FONT_SONG, size=9.5,
                                  bold=(i % 2 == 1), color="#374151" if i % 2 == 0 else f"#{hcolor}")
                else:
                    c_p = cell.add_paragraph()
                    c_p.paragraph_format.left_indent  = Cm(0.5)
                    c_p.paragraph_format.space_before = Pt(1)
                    c_p.paragraph_format.space_after  = Pt(2)
                    parts = re.split(r'\*\*(.+?)\*\*', line)
                    first = True
                    for i, part in enumerate(parts):
                        if not part:
                            continue
                        run_c = c_p.add_run(part)
                        first = False
                        _run_font(run_c, name=FONT_SONG, size=9.5,
                                  bold=(i % 2 == 1), color="#374151" if i % 2 == 0 else f"#{hcolor}")

    # ── 底部页脚 ────────────────────────────────────────────────
    footer_p = cell.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(6)
    footer_p.paragraph_format.space_after  = Pt(4)
    run_f = footer_p.add_run(f"✦ 教学知识卡片  ·  {datetime.now().strftime('%Y.%m.%d')}")
    _run_font(run_f, name=FONT_KAITI, size=8, color="#9CA3AF")


def generate_docx(cards: list[dict], output_path: str) -> Path:
    """生成精美知识卡片 Word 文档"""
    if not HAS_DOCX:
        print("[错误] 未安装 python-docx，请运行：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)

    # 文档大标题
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("📚 教学知识卡片")
    _run_font(run, name=FONT_HEI, size=20, bold=True, color="#1E3A8A")
    title_p.paragraph_format.space_after = Pt(4)

    # 副标题
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub_p.add_run(f"Teaching Knowledge Cards  ·  {datetime.now().strftime('%Y 年 %m 月 %d 日')}")
    _run_font(run_s, name=FONT_KAITI, size=9, color="#6B7280")
    sub_p.paragraph_format.space_after = Pt(14)

    for i, card in enumerate(cards):
        _add_card_to_doc(doc, card)

        if i < len(cards) - 1:
            sep_p = doc.add_paragraph()
            sep_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_sep = sep_p.add_run("· · · · · · ✂ · · · · · · （裁剪线）· · · · · · ✂ · · · · · ·")
            _run_font(run_sep, size=8, color="#D1D5DB")
            sep_p.paragraph_format.space_before = Pt(10)
            sep_p.paragraph_format.space_after  = Pt(10)

            if (i + 1) % 2 == 0:
                doc.add_page_break()

    out_path = Path(output_path)
    doc.save(str(out_path))
    print(f"[✓] 知识卡片 Word 已保存：{out_path}")
    return out_path


# ══════════════════════════════════════════════════════════════════
#   主入口
# ══════════════════════════════════════════════════════════════════

def main():
    import argparse
    parser = argparse.ArgumentParser(description="教学知识卡片生成器 v2.0")
    parser.add_argument('input',  nargs='?', help='Markdown 文件路径')
    parser.add_argument('output', nargs='?', help='输出文件前缀（不含扩展名）')
    parser.add_argument('--text', help='直接传入 Markdown 文本')
    parser.add_argument('--html', action='store_true', default=True,  help='生成 HTML（默认开启）')
    parser.add_argument('--no-html', dest='html', action='store_false', help='禁用 HTML 输出')
    parser.add_argument('--docx', action='store_true', default=True,  help='生成 Word（默认开启）')
    parser.add_argument('--no-docx', dest='docx', action='store_false', help='禁用 Word 输出')
    parser.add_argument('--pdf',  action='store_true', help='同时导出 PDF（需 Word 输出）')
    args = parser.parse_args()

    if args.text:
        md_text = args.text
        output_prefix = args.output or "知识卡片"
    elif args.input:
        md_text = Path(args.input).read_text(encoding='utf-8')
        output_prefix = args.output or Path(args.input).stem
    else:
        print("用法：python knowledge_card_generator.py <input.md> [output_prefix]")
        print("      python knowledge_card_generator.py --text '<markdown>' [output_prefix]")
        sys.exit(1)

    cards = parse_cards(md_text)
    if not cards:
        print("[警告] 未解析到任何知识卡片，请检查 Markdown 格式。")
        sys.exit(1)

    print(f"[✓] 解析到 {len(cards)} 张知识卡片")

    if args.html:
        generate_html(cards, f"{output_prefix}.html")

    docx_path = None
    if args.docx:
        docx_path = generate_docx(cards, f"{output_prefix}.docx")

    if args.pdf and docx_path:
        try:
            pdf_path = Path(f"{output_prefix}.pdf")
            subprocess.run(
                ['python', 'scripts/export_pdf.py', '--input', str(docx_path),
                 '--output', str(pdf_path)],
                check=True
            )
        except Exception as e:
            print(f"[提示] PDF 导出失败：{e}")
            print("       请用浏览器打开 HTML 文件，选择 [打印 - 另存为PDF] 获得最佳效果")


if __name__ == '__main__':
    main()
