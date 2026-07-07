#!/usr/bin/env python3
"""
Markdown → Word 转换工具
接收 Markdown 文本，输出格式化的 .docx 文件。
支持：标题层级、表格、粗体、列表、代码块、分隔线。

用法：
  python scripts/md_to_docx.py <input.md> [output.docx]
  python scripts/md_to_docx.py --text "<markdown内容>" [output.docx]
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

MARGIN = Cm(2.5)

FONT_SONG = "宋体"
FONT_HEI = "黑体"
FONT_KAITI = "楷体"


def _set_cell_shading(cell, color_hex):
    """给单元格设置底色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _run_font(run, font_name=FONT_SONG, font_size=12, bold=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _add_heading(doc, text, level):
    """添加标题——使用 Word 内置标题样式，并设置中文字体"""
    h = doc.add_heading(text.strip(), level=min(level, 5))
    for run in h.runs:
        run.font.name = FONT_HEI
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEI)
    return h


def parse_inline(text):
    """解析行内标记：**粗体**"""
    parts = re.split(r"(\*\*.*?\*\*)", text)
    runs = []
    for p in parts:
        if p.startswith("**") and p.endswith("**"):
            runs.append(("bold", p[2:-2]))
        else:
            runs.append(("normal", p))
    return runs


def add_formatted_paragraph(doc, text, font_size=12, bold=False, alignment=None,
                            space_after=Pt(4), first_line_indent=None):
    """添加带格式的段落（支持**粗体**内联标记）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(22)
    if first_line_indent:
        p.paragraph_format.first_line_indent = first_line_indent
    if alignment:
        p.alignment = alignment

    runs = parse_inline(text)
    for kind, txt in runs:
        run = p.add_run(txt)
        is_bold = bold or (kind == "bold")
        _run_font(run, font_size=font_size, bold=is_bold)

    return p


def add_formatted_paragraphs(doc, text, font_size=12, bold=False, alignment=None,
                             space_after=Pt(4), first_line_indent=None):
    """处理文本中的<br>换行，分隔成多个段落"""
    for line in text.split("\\n"):
        line = line.strip()
        if not line:
            continue
        add_formatted_paragraph(doc, line, font_size, bold, alignment,
                                space_after, first_line_indent)


def add_table_from_md(doc, md_table_text, header_bg="D9E2F3"):
    """
    将 Markdown 表格（以 | 开头的行）转换为 Word 表格
    """
    lines = [l.strip() for l in md_table_text.split("\n") if l.strip() and l.startswith("|")]
    if len(lines) < 2:
        return

    # 跳过分隔行 (|---|---|)
    data_lines = [l for l in lines if not re.match(r"^\|[\s\-:]+\|$", l)]

    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)

    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            font_size = 10
            is_bold = (i == 0)
            _run_font(run, font_size=font_size, bold=is_bold)

            if i == 0:
                _set_cell_shading(cell, header_bg)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def md_to_docx(md_text, output_path, title=None):
    """
    核心转换函数：将 Markdown 文本转为格式化的 Word 文档
    """
    doc = Document()

    # 页面设置
    for section in doc.sections:
        section.top_margin = MARGIN
        section.bottom_margin = MARGIN
        section.left_margin = MARGIN
        section.right_margin = MARGIN

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = FONT_SONG
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_SONG)
    style.paragraph_format.line_spacing = Pt(22)

    lines = md_text.split("\n")
    i = 0
    buf_table = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 空行 → 刷新表格缓冲
        if not stripped:
            if buf_table:
                add_table_from_md(doc, "\n".join(buf_table))
                buf_table = []
            i += 1
            continue

        # 表格行
        if stripped.startswith("|"):
            buf_table.append(stripped)
            i += 1
            continue

        # 刷新表格缓冲
        if buf_table:
            add_table_from_md(doc, "\n".join(buf_table))
            buf_table = []

        # 分隔线 ---
        if re.match(r"^---+$", stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run("─" * 50)
            _run_font(run, font_size=8, color=(180, 180, 180))
            i += 1
            continue

        # 标题
        header_match = re.match(r"^(#{1,4})\s+(.+)$", stripped)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            _add_heading(doc, text, level)
            i += 1
            continue

        # 无序列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            add_formatted_paragraph(doc, f"  {text}", font_size=12,
                                    space_after=Pt(2))
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^\d+[\.\、]\s*(.+)$", stripped)
        if ol_match:
            text = ol_match.group(1)
            add_formatted_paragraph(doc, f"  {stripped}", font_size=12,
                                    space_after=Pt(2))
            i += 1
            continue

        # 普通段落
        add_formatted_paragraph(doc, stripped, font_size=12, space_after=Pt(4))
        i += 1

    # 刷新尾部表格
    if buf_table:
        add_table_from_md(doc, "\n".join(buf_table))

    doc.save(output_path)
    return output_path


def convert_file(input_path, output_path=None):
    """将 Markdown 文件转换为 Word 文档"""
    input_path = Path(input_path)
    if not input_path.exists():
        print(f"文件不存在: {input_path}")
        return None

    if output_path is None:
        output_path = input_path.with_suffix(".docx")

    md_text = input_path.read_text(encoding="utf-8")
    return md_to_docx(md_text, str(output_path))


if __name__ == "__main__":
    if len(sys.argv) >= 3 and sys.argv[1] == "--text":
        # 直接传入 markdown 文本
        md_content = sys.argv[2]
        out_path = sys.argv[3] if len(sys.argv) >= 4 else "output.docx"
        result = md_to_docx(md_content, out_path)
        print(f"Word 文档已生成: {result}")

    elif len(sys.argv) >= 2:
        result = convert_file(sys.argv[1], sys.argv[2] if len(sys.argv) >= 3 else None)
        if result:
            print(f"Word 文档已生成: {result}")
    else:
        print(__doc__)
