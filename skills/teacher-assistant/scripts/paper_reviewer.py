#!/usr/bin/env python3
"""
论文/作业智能批阅器
读取用户提供的文档内容（Markdown/纯文本），
生成带评分、批注和改进建议的批阅报告，支持 Word 和 PDF 导出。

用法：
  python scripts/paper_reviewer.py --input <作业文件路径> [--type <批阅类型>] [--score <满分>]
  python scripts/paper_reviewer.py --text "<作业内容>" [--type <批阅类型>]
"""

import re
import sys
import json
import subprocess
from pathlib import Path
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

FONT_SONG  = "宋体"
FONT_HEI   = "黑体"
FONT_KAITI = "楷体"

# ── 评分维度配置 ──────────────────────────────────────────────────────
RUBRICS = {
    "语文作文": {
        "dimensions": [
            {"name": "内容主题",  "weight": 30, "desc": "立意、选材、内容丰富性"},
            {"name": "结构布局",  "weight": 25, "desc": "段落安排、逻辑连贯"},
            {"name": "语言表达",  "weight": 30, "desc": "词汇丰富、句式多样、表达准确"},
            {"name": "书写规范",  "weight": 5,  "desc": "格式、标点、错别字"},
            {"name": "创新亮点",  "weight": 10, "desc": "有独特见解或精彩语句"},
        ]
    },
    "英语作文": {
        "dimensions": [
            {"name": "内容完整性", "weight": 30, "desc": "是否完整回应题目要求"},
            {"name": "词汇运用",   "weight": 25, "desc": "词汇多样性和准确性"},
            {"name": "语法正确",   "weight": 25, "desc": "时态、句型、语法错误"},
            {"name": "语篇连贯",   "weight": 15, "desc": "衔接词、段落过渡"},
            {"name": "格式规范",   "weight": 5,  "desc": "书信/短文格式"},
        ]
    },
    "数学解答": {
        "dimensions": [
            {"name": "解题思路",  "weight": 30, "desc": "方法选择是否正确"},
            {"name": "解题步骤",  "weight": 40, "desc": "步骤完整、逻辑清晰"},
            {"name": "计算准确",  "weight": 20, "desc": "数值计算是否正确"},
            {"name": "书写规范",  "weight": 10, "desc": "格式规范、单位标注"},
        ]
    },
    "实验报告": {
        "dimensions": [
            {"name": "实验目的",  "weight": 10, "desc": "目的表述清晰"},
            {"name": "实验方法",  "weight": 25, "desc": "方案设计合理"},
            {"name": "数据记录",  "weight": 25, "desc": "数据完整、表格清晰"},
            {"name": "数据分析",  "weight": 25, "desc": "分析准确、有误差讨论"},
            {"name": "实验结论",  "weight": 15, "desc": "结论与数据相符"},
        ]
    },
    "学术论文": {
        "dimensions": [
            {"name": "选题价值",   "weight": 20, "desc": "研究问题是否有意义"},
            {"name": "文献综述",   "weight": 20, "desc": "文献充分、综述有条理"},
            {"name": "论证逻辑",   "weight": 30, "desc": "论证严密、数据支撑"},
            {"name": "格式规范",   "weight": 15, "desc": "引用格式、摘要、参考文献"},
            {"name": "创新贡献",   "weight": 15, "desc": "有新见解或新方法"},
        ]
    },
    "综合作业": {
        "dimensions": [
            {"name": "完整性",    "weight": 25, "desc": "按要求完成所有内容"},
            {"name": "准确性",    "weight": 30, "desc": "知识点运用正确"},
            {"name": "表达能力",  "weight": 25, "desc": "语言流畅、条理清晰"},
            {"name": "拓展延伸",  "weight": 20, "desc": "有个人思考或创意"},
        ]
    },
}


def detect_type(content: str, hint: str = '') -> str:
    """根据内容和提示推断作业类型"""
    if hint:
        for k in RUBRICS:
            if k in hint:
                return k
    lower = content[:500].lower()
    if any(w in lower for w in ['experiment', '实验', '数据', '误差', '结论']):
        return '实验报告'
    if any(w in lower for w in ['abstract', '摘要', '参考文献', 'references', '研究']):
        return '学术论文'
    if any(w in lower for w in ['解', '设', '令', '∵', '∴', '证明', '计算']):
        return '数学解答'
    if any(w in lower for w in ['dear', 'sincerely', 'english', 'writing']):
        return '英语作文'
    return '语文作文'


def _grade_label(score_ratio: float) -> str:
    """评级标签"""
    if score_ratio >= 0.9:   return '优秀'
    if score_ratio >= 0.75:  return '良好'
    if score_ratio >= 0.6:   return '合格'
    return '待提升'


def build_review_report(content: str, work_type: str, total_score: int = 100,
                        student_name: str = '匿名', special_focus: str = '') -> str:
    """
    根据作业内容生成 Markdown 格式批阅报告。
    注意：此函数生成报告的「框架结构」，
    实际批阅评语由调用方（AI 对话流）填充，
    这里提供可独立运行的演示版本（随机示例分）。
    """
    import random
    rubric = RUBRICS.get(work_type, RUBRICS['综合作业'])
    dimensions = rubric['dimensions']
    today = date.today().strftime('%Y年%m月%d日')

    lines = []
    lines.append(f"## 【批阅报告】{work_type} — {student_name}")
    lines.append(f"\n批阅时间：{today}  ")
    lines.append(f"批阅类型：{work_type}  ")
    lines.append(f"满分：{total_score} 分  ")
    if special_focus:
        lines.append(f"特别关注：{special_focus}  ")
    lines.append("\n---\n")

    # ── 一、综合评分 ──
    lines.append("### 一、综合评分\n")
    lines.append("| 评分维度 | 权重 | 得分 | 满分 | 评级 |")
    lines.append("|---------|------|------|------|------|")
    total_get = 0
    dim_scores = []
    for d in dimensions:
        dim_max   = round(total_score * d['weight'] / 100)
        # 演示：随机得分（实际使用时 AI 填入真实评分）
        dim_get   = round(dim_max * random.uniform(0.6, 0.95))
        label     = _grade_label(dim_get / dim_max if dim_max else 0)
        lines.append(f"| {d['name']} | {d['weight']}% | {dim_get} | {dim_max} | {label} |")
        total_get += dim_get
        dim_scores.append((d['name'], dim_get, dim_max, d['desc']))
    grade_label = _grade_label(total_get / total_score)
    lines.append(f"| **总分** | 100% | **{total_get}** | **{total_score}** | **{grade_label}** |")
    lines.append("")

    # ── 二、优点梳理 ──
    lines.append("---\n### 二、优点梳理\n")
    lines.append("**✅ 亮点 1**：（请在此处填写具体优点，引用原文片段支撑）\n")
    lines.append("**✅ 亮点 2**：（请在此处填写第二条亮点）\n")
    lines.append("**✅ 亮点 3**：（如有，填写第三条亮点；否则可删除此行）\n")

    # ── 三、问题批注 ──
    lines.append("---\n### 三、问题批注\n")
    lines.append("**❌ 问题 1**（位置：第 X 段 / 第 X 行）：")
    lines.append("> 原文："……"")
    lines.append("- 问题：（描述具体问题）")
    lines.append("- 建议：（改进方向）")
    lines.append("- 参考修改：（示例改写，可选）\n")
    lines.append("**❌ 问题 2**：（如有，按上述格式继续填写）\n")

    # ── 四、分项详细点评 ──
    lines.append("---\n### 四、分项详细点评\n")
    for name, got, mx, desc in dim_scores:
        lines.append(f"#### {name}（{got}/{mx}分）")
        lines.append(f"*评分依据：{desc}*\n")
        lines.append("（在此处填写 2-4 句详细点评，结合原文内容）\n")

    # ── 五、改进建议 ──
    lines.append("---\n### 五、改进建议\n")
    lines.append("**近期（本次修改）**：")
    lines.append("1. （具体可操作的修改建议，可指明段落位置）")
    lines.append("2. （建议 2）\n")
    lines.append("**长期（综合能力提升）**：")
    lines.append("1. （学习方向/习惯培养）")
    lines.append("2. （推荐阅读或练习方向）\n")

    # ── 六、综合评价 ──
    lines.append("---\n### 六、综合评价\n")
    lines.append("（在此处填写 150-200 字综合性评价：先扬后抑，以鼓励为主，")
    lines.append("指出核心问题，结尾展望进步）\n")

    # ── 七、教师寄语 ──
    lines.append("---\n### 七、教师寄语\n")
    lines.append("（简短温暖的鼓励性文字，1-3 句）\n")
    lines.append("---\n")
    lines.append("*本批阅报告由 AI 辅助生成，建议教师审核后反馈给学生。*")

    return '\n'.join(lines)


def _set_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _run_font(run, name=FONT_SONG, size=11, bold=False, color=None, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    try:
        run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    except Exception:
        pass
    if color:
        r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def md_to_review_docx(md_text: str, output_path: str,
                      student_name: str = '匿名', work_type: str = '综合作业') -> Path:
    """将 Markdown 批阅报告转换为格式化 Word 文档"""
    if not HAS_DOCX:
        print("[错误] 未安装 python-docx，请运行：pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document()

    # A4 页面
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    lines = md_text.split('\n')
    in_table = False
    table_rows = []
    table_obj = None

    def flush_table():
        nonlocal in_table, table_rows, table_obj
        if not table_rows:
            in_table = False
            return
        # 去掉分隔行
        data_rows = [r for r in table_rows if not re.match(r'^\|[-\s|:]+\|$', r.strip())]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        cols = len(re.split(r'\s*\|\s*', data_rows[0].strip('|')))
        tbl = doc.add_table(rows=len(data_rows), cols=cols)
        tbl.style = 'Table Grid'
        for ri, row_str in enumerate(data_rows):
            cells_text = re.split(r'\s*\|\s*', row_str.strip('|'))
            for ci, ct in enumerate(cells_text):
                if ci >= cols:
                    break
                cell = tbl.cell(ri, ci)
                ct_clean = re.sub(r'\*\*(.+?)\*\*', r'\1', ct.strip())
                cell.text = ct_clean
                for run in cell.paragraphs[0].runs:
                    is_header = ri == 0
                    _run_font(run, size=10, bold=is_header,
                              name=FONT_HEI if is_header else FONT_SONG)
                if ri == 0:
                    _set_shading(cell, "BBDEFB")
        in_table = False
        table_rows = []

    for line in lines:
        stripped = line.rstrip()

        # 表格行
        if stripped.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(stripped)
            continue
        elif in_table:
            flush_table()

        # 标题行
        h_match = re.match(r'^(#{1,4})\s+(.+)', stripped)
        if h_match:
            level = len(h_match.group(1))
            text  = h_match.group(2).strip()
            p = doc.add_heading(text, level=min(level, 4))
            for run in p.runs:
                _run_font(run, name=FONT_HEI,
                          size={1: 16, 2: 14, 3: 12, 4: 11}.get(level, 11),
                          bold=True, color="1565C0" if level >= 2 else None)
            continue

        # 引用行
        if stripped.startswith('>'):
            text = stripped.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            run = p.add_run(text)
            _run_font(run, name=FONT_KAITI, size=10, italic=True, color="546E7A")
            continue

        # 列表行
        list_match = re.match(r'^(\s*)([-*•]|\d+\.)\s+(.+)', stripped)
        if list_match:
            indent = len(list_match.group(1))
            text   = list_match.group(3).strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(0.5 + indent * 0.1)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            # 解析粗体
            parts = re.split(r'\*\*(.+?)\*\*', text)
            bullet_run = p.add_run("• ")
            _run_font(bullet_run, size=10, color="1565C0")
            for i, part in enumerate(parts):
                if not part:
                    continue
                run_p = p.add_run(part)
                _run_font(run_p, size=10, bold=(i % 2 == 1))
            continue

        # 分割线
        if re.match(r'^---+$', stripped):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(
                f'<w:pBdr {nsdecls("w")}>'
                f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="BBDEFB"/>'
                f'</w:pBdr>'
            )
            pPr.append(pBdr)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            continue

        # 斜体注释行（*...*）
        italic_match = re.match(r'^\*(.+)\*$', stripped)
        if italic_match:
            p = doc.add_paragraph()
            run = p.add_run(italic_match.group(1))
            _run_font(run, name=FONT_KAITI, size=9, italic=True, color="78909C")
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            continue

        # 普通文字行（含粗体解析）
        if not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(3)
        parts = re.split(r'\*\*(.+?)\*\*', stripped)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            _run_font(run, size=11, bold=(i % 2 == 1))

    if in_table:
        flush_table()

    out_path = Path(output_path)
    doc.save(str(out_path))
    print(f"[✓] 批阅报告 Word 已保存：{out_path}")
    return out_path


def main():
    import argparse
    parser = argparse.ArgumentParser(description="论文/作业智能批阅器")
    parser.add_argument('--input',   help='作业文件路径（.txt/.md）')
    parser.add_argument('--text',    help='直接传入作业内容文本')
    parser.add_argument('--output',  help='输出文件前缀（不含扩展名）')
    parser.add_argument('--type',    help='批阅类型：语文作文/英语作文/数学解答/实验报告/学术论文/综合作业',
                        default='')
    parser.add_argument('--score',   type=int, default=100, help='满分分值（默认100）')
    parser.add_argument('--name',    default='匿名', help='学生姓名')
    parser.add_argument('--focus',   default='', help='特别关注点')
    parser.add_argument('--pdf',     action='store_true', help='同时导出 PDF')
    args = parser.parse_args()

    if args.text:
        content = args.text
        output_prefix = args.output or f"批阅报告_{args.name}"
    elif args.input:
        content = Path(args.input).read_text(encoding='utf-8')
        output_prefix = args.output or f"批阅报告_{Path(args.input).stem}"
    else:
        print("用法：python paper_reviewer.py --input <作业文件> 或 --text <内容>")
        sys.exit(1)

    work_type = detect_type(content, args.type) if not args.type else args.type
    print(f"[✓] 识别批阅类型：{work_type}")

    md_report = build_review_report(
        content, work_type, args.score, args.name, args.focus
    )

    print("\n" + "="*60)
    print(md_report[:800] + "\n…（已截断，完整内容见文件）")
    print("="*60 + "\n")

    docx_path = md_to_review_docx(md_report, f"{output_prefix}.docx",
                                   args.name, work_type)

    if args.pdf:
        try:
            pdf_path = Path(f"{output_prefix}.pdf")
            subprocess.run(
                ['python', 'scripts/export_pdf.py', '--input', str(docx_path),
                 '--output', str(pdf_path)],
                check=True
            )
        except Exception as e:
            print(f"[提示] PDF 导出失败，请手动运行 scripts/export_pdf.py：{e}")


if __name__ == '__main__':
    main()
